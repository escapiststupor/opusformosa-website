from docx import Document

PATH = "/Users/pyen/OpusFormosa/festival_planning/logistics/Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"
document = Document(PATH)

for table_index, activity in [(5, "Schumann: Piano Quartet Op.47"), (10, "Mahler: Piano Quartet in A minor"), (11, "Fauré: Piano Quartet No.1 Op.15")]:
    print(f"TABLE {table_index}")
    for row in document.tables[table_index].rows[1:]:
        if row.cells[2].text == activity:
            values = []
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            values.append((run.text, run.font.size.pt if run.font.size else None))
            print(row.cells[0].text, values)
