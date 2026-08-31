from pathlib import Path

from docx import Document


path = Path("/Users/pyen/OpusFormosa/festival_planning/logistics/Sheraton_Front_Desk.docx")
document = Document(path)

# The front desk needs a simple, consistent operating rule: leave Sheraton
# 25 minutes before any rehearsal at a CHR room.
for paragraph in document.paragraphs:
    if paragraph.text.startswith("下列時間將有車輛在飯店門口等候"):
        paragraph.text = (
            "下列時間將有車輛在飯店門口等候。前往 CHR 排練室時，請於排練開始前 25 分鐘由飯店出發；"
            "表列時間即為出發時間。若有需要，請前台協助通知音樂家。"
        )
        break

times = {
    ("9月3日", "09:30"): "09:35",
    ("9月7日", "09:30"): "09:35",
    ("9月7日", "12:40"): "12:35",
    ("9月8日", "09:30"): "09:35",
    ("9月8日", "12:40"): "12:35",
    ("9月8日", "19:00"): "18:35",
    ("9月9日", "12:40"): "12:35",
    ("9月11日", "09:30"): "09:35",
    ("9月11日", "12:40"): "12:35",
    ("9月13日", "19:00"): "19:05",
}
cars = document.tables[2]
for row in cars.rows[1:]:
    date, time, _, destination, _ = (cell.text for cell in row.cells)
    if destination.startswith("CHR"):
        row.cells[1].text = times[(date, time)]

document.save(path)
