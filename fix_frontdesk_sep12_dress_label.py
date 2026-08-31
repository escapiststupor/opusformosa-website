from pathlib import Path

from docx import Document


SOURCE = Path('/private/tmp/sep12_dress_update/Sheraton_Front_Desk.docx')
OUTPUT = Path('/private/tmp/sep12_dress_update/Sheraton_Front_Desk.docx')


doc = Document(SOURCE)
changes = 0
for table in doc.tables:
    for row in table.rows:
        if not any('Sep 12' in cell.text for cell in row.cells):
            continue
        for cell in row.cells:
            if 'Dress 09:00' in cell.text:
                cell.paragraphs[0].text = cell.text.replace('Dress 09:00', 'Dress 09:30')
                changes += 1
if changes != 1:
    raise RuntimeError(f'Expected one Sep 12 dress label change, found {changes}')
doc.save(OUTPUT)
