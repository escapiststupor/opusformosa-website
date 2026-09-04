from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PATH = Path("/Users/pyen/OpusFormosa/festival_planning/logistics/Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_cell(cell, size=7.5, bold=False, color=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "PingFang TC"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def add_row_after(table, source_row):
    new_row = table.add_row()
    source_row._tr.addnext(new_row._tr)
    return new_row


document = Document(PATH)

# Keep the daily run sheets aligned with the booked / pending tickets.
daily_9 = document.tables[10]
for row in daily_9.rows[1:]:
    if "高鐵前往高雄" in row.cells[2].text:
        row.cells[0].text = "10:31"
        row.cells[1].text = "Edgar Moreau; Jinjoo Cho; Steven Lin"
        row.cells[2].text = "高鐵：台北→左營，車次 121（10:31–12:05）"
        row.cells[3].text = "高鐵"
        row.cells[4].text = "4 席待付款；Edgar 2 席含大提琴。"
        return_row = add_row_after(daily_9, row)
        for cell, value in zip(return_row.cells, [
            "22:10", "Edgar Moreau; Jinjoo Cho; Steven Lin",
            "高鐵：左營→台北，車次 294（22:10–23:49）", "高鐵",
            "4 席待付款；座位分配見「高鐵票券／座位管理」。",
        ]):
            cell.text = value
        break

daily_14 = document.tables[15]
for row in daily_14.rows[1:]:
    if "高鐵前往台中" in row.cells[2].text:
        row.cells[0].text = "11:31"
        row.cells[1].text = "Edgar Moreau; Boris Borgolotto; Steven Lin"
        row.cells[2].text = "高鐵：台北→台中，車次 125（11:31–12:18）"
        row.cells[3].text = "高鐵"
        row.cells[4].text = "4 席待付款；抵達距 13:00 彩排偏緊，建議改早班。"
        return_row = add_row_after(daily_14, row)
        for cell, value in zip(return_row.cells, [
            "22:40", "Edgar Moreau; Boris Borgolotto; Steven Lin",
            "高鐵：台中→台北，車次 862（22:40–23:44）", "高鐵",
            "4 席待付款；座位分配見「高鐵票券／座位管理」。",
        ]):
            cell.text = value
        break

daily_16 = document.tables[17]
for row in daily_16.rows[1:]:
    if "高鐵前往台中" in row.cells[2].text:
        row.cells[0].text = "11:31"
        row.cells[1].text = "Edgar Moreau; Steven Lin"
        row.cells[2].text = "高鐵：台北→台中，車次 125（11:31–12:18）"
        row.cells[3].text = "高鐵"
        row.cells[4].text = "3 席待付款；抵達距 13:00 彩排偏緊，建議改早班。"
        return_row = add_row_after(daily_16, row)
        for cell, value in zip(return_row.cells, [
            "22:40", "Edgar Moreau; Steven Lin",
            "高鐵：台中→台北，車次 862（22:40–23:44）", "高鐵",
            "3 席待付款；座位分配見「高鐵票券／座位管理」。",
        ]):
            cell.text = value
        break

for paragraph in document.paragraphs:
    if paragraph.text == "9月9日，高雄：Jinjoo Cho、Edgar Moreau。":
        paragraph.text = "9月9日，高雄：Edgar Moreau、Jinjoo Cho、Steven Lin。"

# Add a compact ticket-control section immediately before the personnel index.
anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text == "人員索引")
heading = document.add_paragraph("高鐵票券／座位管理")
heading.style = "Heading 1"
note = document.add_paragraph(
    "票券截圖目前均顯示「未付款」。請依 App 顯示的付款期限完成付款；9月9日訂位代號 06478842 的付款期限為 9月3日。"
)
note.style = "normal"
note2 = document.add_paragraph(
    "注意：9月9日的車次 117（09:31）與車次 121（10:31）為重複南下選項，僅保留車次 121；請勿兩班皆付款。"
)
note2.style = "normal"

table = document.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.autofit = False
headers = ["日期", "行程／車次", "席次與人員", "座位", "狀態／執行事項"]
for cell, value in zip(table.rows[0].cells, headers):
    cell.text = value
    shade(cell, "1F4E78")
    style_cell(cell, size=7.5, bold=True)
widths = [0.55, 1.45, 1.35, 1.35, 2.25]
for row in table.rows:
    for cell, width in zip(row.cells, widths):
        cell.width = Inches(width)

rows = [
    ("9/9", "台北→左營\n121｜10:31–12:05", "4 席\nEdgar×2（含大提琴）\nJinjoo、Steven", "截圖未完整顯示\n付款前確認 Edgar 兩席相鄰", "待付款。保留此班；117（09:31）勿再付款。"),
    ("9/9", "左營→台北\n294｜22:10–23:49", "4 席\nEdgar×2、Jinjoo、Steven", "3 車：4D+4E（Edgar）\n5D（Jinjoo）、5E（Steven）", "待付款。"),
    ("9/13", "台北⇄台中\n私人活動", "尚待訂：每程 3 席\nEdgar×2、Steven", "—", "此日期的往返票未見於截圖；須補訂。"),
    ("9/14", "台北→台中\n125｜11:31–12:18", "4 席\nEdgar×2、Boris、Steven", "5 車：11D+11E（Edgar）\n12D（Boris）、12E（Steven）", "待付款。距 13:00 彩排僅 42 分鐘；建議改早班。"),
    ("9/14", "台中→台北\n862｜22:40–23:44", "4 席\nEdgar×2、Boris、Steven", "2 車：6D+6E（Edgar）\n7D（Boris）、7E（Steven）", "待付款。"),
    ("9/16", "台北→台中\n125｜11:31–12:18", "3 席\nEdgar×2、Steven", "5 車：17A+17B（Edgar）\n17C（Steven）", "待付款。距 13:00 彩排僅 42 分鐘；建議改早班。"),
    ("9/16", "台中→台北\n862｜22:40–23:44", "3 席\nEdgar×2、Steven", "4 車：9A+9B（Edgar）\n9C（Steven）", "待付款。"),
]
for values in rows:
    row = table.add_row()
    for cell, value, width in zip(row.cells, values, widths):
        cell.text = value
        cell.width = Inches(width)
        style_cell(cell)

# Move the new section before the personnel index, retaining the handbook flow.
anchor._p.addprevious(heading._p)
anchor._p.addprevious(note._p)
anchor._p.addprevious(note2._p)
anchor._p.addprevious(table._tbl)

document.save(PATH)
