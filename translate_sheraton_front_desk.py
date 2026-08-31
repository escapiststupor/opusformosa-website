from pathlib import Path

from docx import Document


path = Path("/Users/pyen/OpusFormosa/festival_planning/logistics/Sheraton_Front_Desk.docx")
document = Document(path)


def set_paragraph(paragraph, text):
    paragraph.text = text


paragraphs = {
    "Sheraton Grand Taipei — Guest Logistics": "台北喜來登大飯店 — 貴賓接待與交通安排",
    "Opus Formosa 2026  |  Sep 2–16  |  Confidential": "福爾摩沙藝響 2026 ｜9月2日至16日｜機密",
    "Festival Guests": "音樂節貴賓",
    "Key Contacts": "主要聯絡人",
    "Car Pickup Schedule": "接送車輛時刻表",
    "A car will be ready at the hotel entrance at the times listed below. Please ensure musicians are notified by front desk if needed.": "下列時間將有車輛在飯店門口等候。若有需要，請前台協助通知音樂家。",
    "General Notes": "一般注意事項",
    "All airport arrivals are handled by dedicated chauffeur service — no action needed from Sheraton.": "所有機場抵達均由專屬接機車輛處理，喜來登前台無需另外安排。",
    "If a musician cannot reach their car or driver, they should contact Phyllis Canglah M. or Ethan (Yiting) by WhatsApp.": "若音樂家找不到接送車輛或司機，請協助其透過 WhatsApp 聯絡 Phyllis Canglah M. 或 Ethan（Yiting）。",
    "CHR1–4 = National Concert Hall backstage rehearsal rooms (10 min drive).": "CHR1–4＝國家音樂廳後台排練室（車程約10分鐘）。",
    "Sep 9: Jinjoo Cho and Edgar Moreau travel to Kaohsiung by HSR for the evening concert. Return late night.": "9月9日：Jinjoo Cho 與 Edgar Moreau 搭乘高鐵前往高雄參加晚間音樂會，深夜返回。",
    "Sep 12 (Saturday): Dress rehearsal begins at 09:30.": "9月12日（星期六）：彩排於09:30開始。",
    "Key locations & access": "重要地點與出入方式",
    "All CHR rehearsal rooms (CHR1 / CHR2 / CHR3 / CHR4) — National Concert Hall Backstage, Basement Floor, No. 21-1, Zhongshan S. Rd., Zhongzheng Dist., Taipei City 100012, Taiwan. Use the staff entrance: https://maps.app.goo.gl/kxPsa1KHaKpFppgi9. Please bring your ID card or passport and show it to the guard before entering.": "所有 CHR 排練室（CHR1／CHR2／CHR3／CHR4）位於國家音樂廳後台地下樓，地址：台北市中正區中山南路21-1號。請由工作人員入口進出：https://maps.app.goo.gl/kxPsa1KHaKpFppgi9。進入前須攜帶身分證或護照，並向警衛出示。",
    "National Concert Hall / National Recital Hall — No. 21-1, Zhongshan S. Rd., Zhongzheng Dist., Taipei City 100012, Taiwan.": "國家音樂廳／國家演奏廳：台北市中正區中山南路21-1號。",
    "National Taichung Theater — 101, Huilai Rd., Sec. 2, Xitun District, Taichung City 407025, Taiwan.": "臺中國家歌劇院：台中市西屯區惠來路二段101號。",
    "National Kaohsiung Center for the Arts (Weiwuying) — No. 1, Sanduo 1st Rd., Fengshan Dist., Kaohsiung City 830043, Taiwan.": "衛武營國家藝術文化中心：高雄市鳳山區三多一路1號。",
    "Sponsor's Dinner — 1F., No. 70-1, Chengde Rd., Sec. 1, Datong Dist., Taipei City, Taiwan.": "贊助人晚宴：台北市大同區承德路一段70-1號1樓。",
    "HSR travel companions": "高鐵同行名單",
    "Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.": "9月9日，高雄：Jinjoo Cho、Edgar Moreau。",
    "Sep 13, Taichung: Edgar Moreau and Steven Lin.": "9月13日，台中：Edgar Moreau、Steven Lin。",
    "Sep 14, Taichung: Edgar Moreau, Boris Borgolotto, and Steven Lin.": "9月14日，台中：Edgar Moreau、Boris Borgolotto、Steven Lin。",
    "Sep 16, Taichung: Edgar Moreau and Steven Lin.": "9月16日，台中：Edgar Moreau、Steven Lin。",
}
for paragraph in document.paragraphs:
    if paragraph.text in paragraphs:
        set_paragraph(paragraph, paragraphs[paragraph.text])

# Keep the QR image intact while translating its label.
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = run.text.replace("WhatsApp local contact", "WhatsApp 當地聯絡人")


guests = document.tables[1]
for cell, value in zip(guests.rows[0].cells, ["姓名", "抵達", "離境"]):
    cell.text = value
guest_rows = {
    "Adrien La Marca": (
        "抵達：9月3日 06:50 — BR88，巴黎戴高樂（CDG）→ 台北桃園（TPE），第2航廈",
        "離境：9月13日 06:30 — BR178，台北桃園（TPE），第2航廈 → 大阪關西",
    ),
    "Aimi Kobayashi": (
        "抵達：9月2日 15:50 — NH853，東京羽田（HND）→ 台北松山（TSA），第1航廈",
        "離境：9月6日 13:30 — NH852，台北松山（TSA），第1航廈 → 東京羽田\n不安排送機。",
    ),
    "Brannon Cho": (
        "抵達：9月6日 18:45 — UA871，美聯航 → 台北桃園（TPE），第2航廈",
        "離境：9月13日 14:30 — UA852，美聯航，台北桃園（TPE），第2航廈",
    ),
    "Jinjoo Cho": (
        "抵達：9月7日 21:25 — BR159，長榮航空 → 台北桃園（TPE），第2航廈",
        "離境：9月13日 15:15 — BR160，長榮航空，台北桃園（TPE），第2航廈",
    ),
    "Kyu Yeon Kim": (
        "抵達：9月7日 10:25 — KE5659（韓亞航空 OZ711 執飛），首爾仁川（ICN）→ 台北桃園（TPE），第2航廈",
        "離境：9月13日 13:20 — KE2022，大韓航空，台北桃園（TPE），第1航廈",
    ),
    "Edgar Moreau": (
        "抵達：9月8日 06:50 — BR88，巴黎戴高樂（CDG）→ 台北桃園（TPE），第2航廈\n抵達後安排半日早入住。",
        "離境：9月17日 23:30 — BR87，台北桃園（TPE），第2航廈 → 巴黎戴高樂",
    ),
}
for row in guests.rows[1:]:
    arrival, departure = guest_rows[row.cells[0].text]
    row.cells[1].text = arrival
    row.cells[2].text = departure


cars = document.tables[2]
for cell, value in zip(cars.rows[0].cells, ["日期", "時間", "乘客", "目的地", "備註"]):
    cell.text = value
car_rows = [
    ("9月2日", "待定", "Aimi Kobayashi", "贊助人晚宴會場", "由當地協調人陪同"),
    ("9月3日", "09:30", "Adrien La Marca", "CHR3（國家音樂廳後台）", ""),
    ("9月3日", "14:30", "Aimi Kobayashi", "史坦威中心", "由當地協調人陪同"),
    ("9月4日", "12:40", "Adrien La Marca、Aimi Kobayashi", "國家音樂廳", "彩排"),
    ("9月7日", "10:25", "Kyu Yeon Kim", "桃園機場第2航廈 → 喜來登", "舉牌接機；優先入住及領取後台證"),
    ("9月7日", "09:30", "Adrien La Marca、Brannon Cho", "CHR4（國家音樂廳後台）", ""),
    ("9月7日", "12:40", "Adrien La Marca、Kyu Yeon Kim", "CHR4（國家音樂廳後台）", "Kyu Yeon 優先入住後共乘"),
    ("9月8日", "09:30", "Adrien La Marca、Brannon Cho", "CHR4（國家音樂廳後台）", ""),
    ("9月8日", "12:40", "Edgar Moreau", "CHR2（國家音樂廳後台）", "首日抵達（早上）"),
    ("9月8日", "12:40", "Jinjoo Cho、Brannon Cho、Kyu Yeon Kim", "CHR4（國家音樂廳後台）", ""),
    ("9月8日", "19:00", "Adrien La Marca、Edgar Moreau", "CHR4（國家音樂廳後台）", "德弗札克晚間排練"),
    ("9月9日", "10:00", "Jinjoo Cho、Edgar Moreau", "台北車站（高鐵 → 高雄）", "衛武營音樂會"),
    ("9月9日", "12:40", "Adrien La Marca、Brannon Cho、Kyu Yeon Kim", "CHR4（國家音樂廳後台）", "佛瑞排練"),
    ("9月10日", "12:40", "Adrien、Jinjoo、Brannon、Kyu Yeon、Edgar", "國家演奏廳", "彩排"),
    ("9月11日", "09:30", "Edgar Moreau", "CHR4（國家音樂廳後台）", "布拉姆斯上午排練"),
    ("9月11日", "12:40", "Adrien La Marca、Brannon Cho、Jinjoo Cho、Edgar Moreau", "CHR4", "柴科夫斯基排練"),
    ("9月12日", "08:30", "Adrien、Jinjoo、Brannon、Kyu Yeon、Edgar", "國家演奏廳", "09:30 彩排"),
    ("9月13日", "09:00", "Edgar Moreau", "台北車站（高鐵 → 台中）", "與 Steven 的私人活動"),
    ("9月13日", "19:00", "Edgar Moreau", "CHR4（國家音樂廳後台）", "布拉姆斯晚間排練"),
]
for row, values in zip(cars.rows[1:], car_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

for paragraph, text in zip(document.paragraphs[6:12], [
    "Phyllis Canglah M. — WhatsApp：https://wa.me/qr/BBW55DSGXCD5H1\nEthan（Yiting）— WhatsApp：https://wa.me/qr/G7JIT5KAJTLON1",
    "Ethan：Ethan —［電話］",
    "Glandys：Glandys —［電話］",
    "車隊服務：Car Service —［電話］",
    "喜來登：台北喜來登大飯店 — No. 12, Sec. 1, Zhongxiao East Road, Zhongzheng District, Taipei, Taiwan 10049 — +886-2-2321-5511",
    "Phyllis：Phyllis —［電話］",
]):
    set_paragraph(paragraph, text)

# The original blank paragraph contained a manual page break.  Removing it
# keeps the operational notes and the location section together.
document.paragraphs[23].text = ""

document.save(path)
