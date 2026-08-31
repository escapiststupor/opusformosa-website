from copy import deepcopy
from pathlib import Path

from docx import Document


path = Path('/Users/pyen/OpusFormosa/festival_planning/logistics/Kyu_Yeon_Kim.docx')
doc = Document(path)

# Correct the Sep 7 practice-room row in the main schedule table.
schedule = next(
    table for table in doc.tables
    if [cell.text.strip() for cell in table.rows[0].cells][:2] == ['Date', 'Time']
)
schedule_row = next(
    row for row in schedule.rows
    if row.cells[1].text.strip() == '14:00–19:00'
    and 'Practice room' in row.cells[2].text
)
schedule_row.cells[1].text = '16:00–19:00'

# Add the missing Sep 7 reservation line immediately before the Sep 8 line,
# retaining the existing paragraph formatting.
reservation_sep8 = next(
    paragraph for paragraph in doc.paragraphs
    if paragraph.text.strip().startswith('Sep 8: 10:00–12:00 — grand piano')
)
if not any(paragraph.text.strip().startswith('Sep 7:') for paragraph in doc.paragraphs):
    copied = deepcopy(reservation_sep8._p)
    for text_node in copied.iter():
        if text_node.tag.endswith('}t'):
            if text_node.text == 'Sep 8: ':
                text_node.text = 'Sep 7: '
            elif text_node.text == '10:00–12:00 — grand piano':
                text_node.text = '16:00–19:00 — grand piano'
    reservation_sep8._p.addprevious(copied)

doc.save(path)
print(path)
