from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')


def row_values(row):
    return [cell.text.strip() for cell in row.cells]


def set_row(row, values):
    for cell, value in zip(row.cells, values):
        cell.text = value


def fix_kyu_personal():
    path = ROOT / 'Kyu_Yeon_Kim.docx'
    doc = Document(path)
    table = next(t for t in doc.tables if row_values(t.rows[0])[:4] == ['Date', 'Time', 'Programme / Event', 'Venue'])
    sep8_index = next(i for i, row in enumerate(table.rows) if row_values(row)[0] == 'Sep 8')
    incorrect_index = next(i for i, row in enumerate(table.rows) if row_values(row)[1] == '16:00–19:00')
    incorrect = table.rows[incorrect_index]
    copied_row = deepcopy(incorrect._tr)
    table._tbl.insert(sep8_index, copied_row)
    inserted = table.rows[sep8_index]
    set_row(inserted, ['', '14:00–19:00', 'Practice room — grand piano', 'Kawaii Studio'])
    table._tbl.remove(incorrect._tr)
    doc.save(path)


def fix_handbook():
    path = ROOT / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx'
    doc = Document(path)
    replacements = 0
    for table in doc.tables:
        for row in table.rows:
            values = row_values(row)
            if values and values[0] == 'Sep 7 16:00-19:00' and len(values) >= 5 and values[1] == 'Kyu Yeon Kim':
                row.cells[0].text = 'Sep 7 14:00-19:00'
                replacements += 1
    if replacements != 1:
        raise RuntimeError(f'Expected one handbook Kyu practice row, found {replacements}')
    doc.save(path)


fix_kyu_personal()
fix_handbook()
print('Updated Kyu personal schedule and Da Zong Guan handbook.')
