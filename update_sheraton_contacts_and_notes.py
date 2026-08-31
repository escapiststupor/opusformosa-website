from pathlib import Path

from docx import Document


path = Path("/Users/pyen/OpusFormosa/festival_planning/logistics/Sheraton_Front_Desk.docx")
document = Document(path)

for paragraph in document.paragraphs:
    if paragraph.text.startswith("Phyllis Canglah M."):
        paragraph.text = (
            "Phyllis Canglah M.｜電話：0920-921-779｜WhatsApp：https://wa.me/qr/BBW55DSGXCD5H1\n"
            "Ethan（Yiting）｜WhatsApp：https://wa.me/qr/G7JIT5KAJTLON1\n"
            "Glandys｜WhatsApp：https://wa.me/qr/DOER7EC3KLVSG1"
        )
    elif paragraph.text.startswith("Ethan："):
        paragraph.text = "Ethan（Yiting）：WhatsApp：https://wa.me/qr/G7JIT5KAJTLON1"
    elif paragraph.text.startswith("Glandys："):
        paragraph.text = "Glandys：WhatsApp：https://wa.me/qr/DOER7EC3KLVSG1"
    elif paragraph.text.startswith("Phyllis："):
        paragraph.text = "Phyllis Canglah M.：0920-921-779"
    elif paragraph.text.startswith("所有機場抵達"):
        paragraph.text = "機場接機由主辦方另行處理；喜來登前台不需提供或安排接機服務。"

cars = document.tables[2]
for row in cars.rows[1:]:
    date, time, passengers, destination, notes = (cell.text for cell in row.cells)
    if date == "9月2日" and passengers == "Aimi Kobayashi":
        row.cells[4].text = "當地協調人：Glandys（陪同）"
    if date == "9月8日" and time == "12:35" and passengers == "Edgar Moreau":
        row.cells[4].text = "首日抵達（早上）；中午由當地協調人至飯店問候"

document.save(path)
