from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
QR = Path('/var/folders/64/y4m8r3fj3f5020hryd4sq5mh0000gp/T/codex-clipboard-ee427d0f-bf5a-493e-acf7-eef4f297e04a.png')
CONTACT_TEXT = 'Glandys — WhatsApp local contact\nhttps://wa.me/qr/DOER7EC3KLVSG1'


def has_contact(doc):
    return 'DOER7EC3KLVSG1' in '\n'.join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )


def add_label_and_qr(cell, width=0.55):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = paragraph.add_run('Glandys\nWhatsApp local contact\n')
    label.bold = True
    label.font.size = Pt(8)
    paragraph.add_run('https://wa.me/qr/DOER7EC3KLVSG1\n').font.size = Pt(6.5)
    picture = paragraph.add_run().add_picture(str(QR), width=Inches(width))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_to_existing_contact_table(doc, table):
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    add_label_and_qr(cell, width=0.52)


def add_new_contact_block(doc):
    block = doc.add_table(rows=1, cols=1)
    block.style = 'Table Grid'
    add_label_and_qr(block.cell(0, 0), width=0.50)
    if doc.tables:
        # The new table is appended; move it directly before the first existing table,
        # preserving the document's title and introductory paragraphs above it.
        first_existing = doc.tables[0]
        if first_existing is not block:
            first_existing._tbl.addprevious(block._tbl)


def find_existing_contact_table(doc):
    for table in doc.tables:
        text = '\n'.join(cell.text for row in table.rows for cell in row.cells)
        if 'Phyllis Canglah M.' in text and 'Ethan (Yiting)' in text:
            # Only the foreign-artist contact cards are a two-column QR card.
            if len(table.columns) == 2 and len(table.rows) == 2:
                return table
    return None


if not QR.exists():
    raise FileNotFoundError(QR)

updated = []
for path in sorted(ROOT.glob('*.docx')):
    doc = Document(path)
    if has_contact(doc):
        continue
    existing = find_existing_contact_table(doc)
    if existing:
        add_to_existing_contact_table(doc, existing)
    else:
        add_new_contact_block(doc)
    doc.save(path)
    updated.append(path.name)

if len(updated) != 19:
    raise RuntimeError(f'Expected 19 logistics DOCX files to update, changed {len(updated)}: {updated}')
print('\n'.join(updated))
