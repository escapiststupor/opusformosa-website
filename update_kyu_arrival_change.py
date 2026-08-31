from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/kyu_arrival_change_update')

ARRIVAL = 'Sep 7, 10:25 — KE5659 (operated by Asiana Airlines OZ711), Seoul Incheon (ICN) → Taipei Taoyuan (TPE), Terminal 2'
ARRIVAL_SHORT = 'Arrival: KE5659 (operated by Asiana Airlines OZ711), ICN → TPE Terminal 2'


def delete_row(row):
    row._tr.getparent().remove(row._tr)


def insert_before(table, target_row, source_row, values):
    new_tr = deepcopy(source_row._tr)
    table._tbl.insert(table._tbl.index(target_row._tr), new_tr)
    new_row = next(row for row in table.rows if row._tr is new_tr)
    for cell, value in zip(new_row.cells, values):
        cell.text = value
    return new_row


def replace_all_text(doc, old, new):
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)


def update_kyu_itinerary():
    doc = Document(SOURCE / 'Kyu_Yeon_Kim.docx')
    for p in doc.paragraphs:
        if p.text.startswith('Arrival:'):
            p.text = 'Arrival:  ' + ARRIVAL
        elif p.text.startswith('Note:'):
            p.text = p.text.replace('Taxi from Sheraton for all rehearsals.', 'On Sep 7, our team will meet her at TPE Terminal 2, prioritise a quick Sheraton check-in and stage-pass collection, then arrange a shared 12:30 taxi with Adrien La Marca to CHR4.')
        elif p.text.startswith('Sep 7: 10:00–12:00'):
            p._element.getparent().remove(p._element)

    schedule = doc.tables[1]
    practice_row = next(row for row in schedule.rows if row.cells[0].text.strip() == 'Sep 7' and row.cells[1].text.strip() == '10:00–12:00')
    delete_row(practice_row)
    first_rehearsal = next(row for row in schedule.rows if row.cells[1].text.strip() == '13:00–15:00' and 'Mahler: Piano Quartet in A minor' in row.cells[2].text)
    template = schedule.rows[1]
    insert_before(schedule, first_rehearsal, template, ['Sep 7', '10:25', 'Airport arrival — KE5659 (operated by Asiana Airlines OZ711)', 'Taipei Taoyuan (TPE), Terminal 2'])
    insert_before(schedule, first_rehearsal, template, ['', '~11:45–12:20', 'Priority hotel check-in; collect stage pass', 'Sheraton Grand Taipei'])
    insert_before(schedule, first_rehearsal, template, ['', '12:30', 'Shared taxi with Adrien La Marca to CHR4', 'Sheraton → CHR4'])
    first_rehearsal.cells[0].text = ''
    return doc


def update_front_desk():
    doc = Document(SOURCE / 'Sheraton_Front_Desk.docx')
    replace_all_text(doc, 'Kyuyeon', 'Kyu Yeon')
    guests = doc.tables[0]
    for row in guests.rows:
        if row.cells[0].text.strip() == 'Kyu Yeon Kim':
            row.cells[1].text = 'Arrival: ' + ARRIVAL
    cars = doc.tables[1]
    template = cars.rows[1]
    first_sep7 = next(row for row in cars.rows if row.cells[0].text.strip() == 'Sep 7')
    insert_before(cars, first_sep7, template, ['Sep 7', '10:25', 'Kyu Yeon Kim', 'TPE Terminal 2 → Sheraton', 'Name-sign pickup; priority check-in and stage pass'])
    for row in cars.rows:
        if row.cells[0].text.strip() == 'Sep 7' and row.cells[1].text.strip() == '12:30':
            row.cells[2].text = 'Adrien La Marca, Kyu Yeon Kim'
            row.cells[4].text = 'Shared car after Kyu Yeon\'s priority check-in'
    return doc


def update_master_schedule():
    doc = Document(SOURCE / 'Da_Zong_Guan_Master_Schedule.docx')
    replace_all_text(doc, 'Kyuyeon', 'Kyu Yeon')
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            all_text = ' | '.join(cell.text for cell in cells)
            if 'Kyu Yeon Kim' in all_text and 'LJ763' in all_text:
                cells[1].text = ARRIVAL
            if 'Airport chauffeur for Jinjoo and Kyu Yeon' in all_text:
                cells[1].text = '10:25 / 21:25'
                cells[2].text = 'Airport pickup: Kyu Yeon arrives 10:25 at TPE Terminal 2 (KE5659 / Asiana OZ711). Arrange priority Sheraton check-in and stage pass; she joins Adrien in the shared 12:30 car to CHR4. Airport chauffeur for Jinjoo at 21:25; confirm Sheraton check-in.'
    return doc


def update_handbook():
    doc = Document(SOURCE / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')
    replace_all_text(doc, 'Sep 7 00:15', 'Sep 7 10:25')
    replace_all_text(doc, '00:15', '10:25')
    replace_all_text(doc, 'Arrival: Sep 7, 00:15 — LJ763, Jin Air → Taipei Taoyuan (TPE), Terminal 1', 'Arrival: ' + ARRIVAL)
    replace_all_text(doc, 'Arrival: LJ763 Jeju to TPE Terminal 1', ARRIVAL_SHORT)
    for table in doc.tables:
        is_sep7_daily_table = 'KE5659' in ' | '.join(cell.text for row in table.rows for cell in row.cells)
        for row in list(table.rows):
            row_text = ' | '.join(cell.text for cell in row.cells)
            if 'Sep 7 10:00-12:00' in row_text and 'Kyu Yeon Kim' in row_text:
                delete_row(row)
                continue
            if is_sep7_daily_table and 'Kyu Yeon Kim' in row_text and 'Practice room: grand piano' in row_text and row.cells[0].text.strip().replace('\n', '').startswith('10:00'):
                delete_row(row)
                continue
            if 'Kyu Yeon Kim' in row_text and 'LJ763' in row_text:
                for cell in row.cells:
                    if 'LJ763' in cell.text:
                        cell.text = 'Arrival: ' + ARRIVAL
            if 'Kyu Yeon Kim' in row_text and 'Airport pickup' in row_text and ('KE5659' in row_text or 'OZ711' in row_text):
                row.cells[-1].text = 'Airport pickup; priority Sheraton check-in and stage pass; shared 12:30 car with Adrien to CHR4'
    return doc


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    update_kyu_itinerary().save(OUTPUT / 'Kyu_Yeon_Kim.docx')
    update_front_desk().save(OUTPUT / 'Sheraton_Front_Desk.docx')
    update_master_schedule().save(OUTPUT / 'Da_Zong_Guan_Master_Schedule.docx')
    update_handbook().save(OUTPUT / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')


if __name__ == '__main__':
    main()
