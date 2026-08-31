from docx import Document
from docx.shared import Pt


PATH = "/Users/pyen/OpusFormosa/festival_planning/logistics/Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"

document = Document(PATH)
for table in document.tables[4:20]:
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "PingFang TC"
                    run.font.size = Pt(7.5)

document.save(PATH)
print("Normalized daily run-sheet row typography.")
