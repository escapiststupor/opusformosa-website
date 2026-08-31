from pathlib import Path
from shutil import copy2

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/opus_flight_details_update')

FLIGHTS = {
    'Aimi Kobayashi': {
        'file': 'Aimi_Kobayashi.docx',
        'arrival': 'Sep 2, 15:50 — NH853, Tokyo Haneda (HND) → Taipei Songshan (TSA), Terminal 1',
        'departure': 'Sep 6, 13:30 — NH852, Taipei Songshan (TSA), Terminal 1 → Tokyo Haneda (HND)',
    },
    'Adrien La Marca': {
        'file': 'Adrien_La_Marca.docx',
        'arrival': 'Sep 3, 06:50 — BR88, Paris CDG → Taipei Taoyuan (TPE), Terminal 2',
        'departure': 'Sep 13, 06:30 — BR178, Taipei Taoyuan (TPE), Terminal 2 → Osaka Kansai',
    },
    'Edgar Moreau': {
        'file': 'Edgar_Moreau.docx',
        'arrival': 'Sep 8, 06:50 — BR88, Paris CDG → Taipei Taoyuan (TPE), Terminal 2',
        'departure': 'Sep 17, 23:30 — BR87, Taipei Taoyuan (TPE), Terminal 2 → Paris CDG',
    },
    'Jinjoo Cho': {
        'file': 'Jinjoo_Cho.docx',
        'arrival': 'Sep 7, 21:25 — BR159, EVA Air → Taipei Taoyuan (TPE), Terminal 2',
        'departure': 'Sep 13, 15:15 — BR160, EVA Air, Taipei Taoyuan (TPE), Terminal 2',
    },
    'Brannon Cho': {
        'file': 'Brannon_Cho.docx',
        'arrival': 'Sep 6, 18:45 — UA871, United → Taipei Taoyuan (TPE), Terminal 2',
        'departure': 'Sep 13, 14:30 — UA852, United, Taipei Taoyuan (TPE), Terminal 2',
    },
    'Kyu Yeon Kim': {
        'file': 'Kyu_Yeon_Kim.docx',
        'arrival': 'Sep 7, 00:15 — LJ763, Jin Air → Taipei Taoyuan (TPE), Terminal 1',
        'departure': 'Sep 13, 13:20 — KE2022, Korean Air, Taipei Taoyuan (TPE), Terminal 1',
    },
    'Boris Borgolotto': {
        'file': 'Boris_Borgolotto.docx',
        'arrival': 'Sep 8, 11:40 — CA185, Air China → Taipei Taoyuan (TPE), Terminal 2',
        'departure': 'Sep 15, 09:10 — CA5516, Air China, Taipei Taoyuan (TPE), Terminal 2',
    },
}


def replace_personal_flight_lines(doc, details):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith('Arrival:'):
            paragraph.text = 'Arrival:  ' + details['arrival']
        elif paragraph.text.startswith('Departure:'):
            paragraph.text = 'Departure:  ' + details['departure']


def replace_table_cell(cell, text):
    cell.text = text


def update_front_desk(path):
    doc = Document(path)
    guest_table = doc.tables[0]
    for row in guest_table.rows[1:]:
        name = row.cells[0].text.strip()
        if name == 'Kyuyeon Kim':
            name = 'Kyu Yeon Kim'
            row.cells[0].text = name
        if name in FLIGHTS:
            replace_table_cell(row.cells[1], 'Arrival: ' + FLIGHTS[name]['arrival'])
            replace_table_cell(row.cells[2], 'Departure: ' + FLIGHTS[name]['departure'])
    return doc


def update_master_schedule(path):
    doc = Document(path)
    doc.add_heading('Airport Flight Details', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for cell, text in zip(table.rows[0].cells, ['Artist', 'Arrival flight', 'Departure flight']):
        cell.text = text
    for name, details in FLIGHTS.items():
        row = table.add_row().cells
        row[0].text = name
        row[1].text = details['arrival']
        row[2].text = details['departure']
    return doc


def update_handbook(path):
    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 4:
                continue
            name = row.cells[1].text.strip()
            if name not in FLIGHTS:
                continue
            when = row.cells[0].text.strip()
            action = row.cells[-1].text.strip()
            if action == 'Airport pickup':
                row.cells[3].text = 'Arrival: ' + FLIGHTS[name]['arrival']
            elif action == 'Airport drop-off':
                row.cells[3].text = 'Departure: ' + FLIGHTS[name]['departure']
            elif when:
                # The first travel-control table uses the action in the final column too.
                text = row.cells[3].text
                if text.startswith('Arrival:'):
                    row.cells[3].text = 'Arrival: ' + FLIGHTS[name]['arrival']
                elif text.startswith('Departure:'):
                    row.cells[3].text = 'Departure: ' + FLIGHTS[name]['departure']
    return doc


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for details in FLIGHTS.values():
        source = SOURCE / details['file']
        doc = Document(source)
        replace_personal_flight_lines(doc, details)
        doc.save(OUTPUT / source.name)

    front = update_front_desk(SOURCE / 'Sheraton_Front_Desk.docx')
    front.save(OUTPUT / 'Sheraton_Front_Desk.docx')

    master = update_master_schedule(SOURCE / 'Da_Zong_Guan_Master_Schedule.docx')
    master.save(OUTPUT / 'Da_Zong_Guan_Master_Schedule.docx')

    handbook = update_handbook(SOURCE / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')
    handbook.save(OUTPUT / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')


if __name__ == '__main__':
    main()
