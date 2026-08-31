from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
KEYS = ('CHR rehearsal rooms', 'CHR1–4:', 'Sheraton:', 'National Concert Hall Backstage')


def paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text


for path in sorted(ROOT.glob('*.docx')):
    matches = [text for text in paragraphs(Document(path)) if any(key in text for key in KEYS)]
    if matches:
        print('---' + path.name + '---')
        print('\n'.join(matches))
