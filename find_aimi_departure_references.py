from pathlib import Path

from docx import Document


ROOT = Path("/Users/pyen/OpusFormosa/festival_planning/logistics")
for path in ROOT.glob("*.docx"):
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    matches = [chunk.replace("\n", " / ") for chunk in chunks if "NH852" in chunk or ("Aimi" in chunk and "送機" in chunk)]
    if matches:
        print(path.name)
        for match in matches:
            print("  ", match)
