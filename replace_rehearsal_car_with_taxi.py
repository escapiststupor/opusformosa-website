from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/taxi_rehearsal_update')
REPLACEMENTS = {
    'Adrien_La_Marca.docx': {
        'Car from Sheraton for all rehearsals.': 'Taxi from Sheraton for all rehearsals.',
    },
    'Aimi_Kobayashi.docx': {
        'Car from Sheraton for the Sep 3 Steinway Center session.': 'Taxi from Sheraton for the Sep 3 Steinway Center session.',
    },
    'Brannon_Cho.docx': {
        'Car from Sheraton for all rehearsals.': 'Taxi from Sheraton for all rehearsals.',
    },
    'Edgar_Moreau.docx': {
        'Car from Sheraton for all rehearsals.': 'Taxi from Sheraton for all rehearsals.',
    },
    'Jinjoo_Cho.docx': {
        'Car from Sheraton for all rehearsals.': 'Taxi from Sheraton for all rehearsals.',
    },
    'Kyu_Yeon_Kim.docx': {
        'Car from Sheraton for all rehearsals.': 'Taxi from Sheraton for all rehearsals.',
    },
}


def replace_in_paragraph(paragraph, replacements):
    original = paragraph.text
    updated = original
    for before, after in replacements.items():
        updated = updated.replace(before, after)
    if updated != original:
        paragraph.text = updated


def update_document(source, output, replacements):
    doc = Document(source)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    doc.save(output)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for filename, replacements in REPLACEMENTS.items():
        update_document(SOURCE / filename, OUTPUT / filename, replacements)


if __name__ == '__main__':
    main()
