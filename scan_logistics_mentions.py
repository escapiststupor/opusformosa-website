from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
KEYS = [
    'CHR1', 'CHR2', 'Sponsor', 'sponsor', 'National Concert Hall',
    'National Recital Hall', 'Weiwuying', 'Taichung',
]


def text_of(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return '\n'.join(parts)


for path in sorted(ROOT.glob('*.docx')):
    found = [key for key in KEYS if key in text_of(Document(path))]
    if found:
        print(f'{path.name} | {", ".join(found)}')
