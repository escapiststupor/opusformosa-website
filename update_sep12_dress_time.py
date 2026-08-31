from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/sep12_dress_update')
TABLE_TARGETS = {
    'Adrien_La_Marca.docx',
    'Boris_Borgolotto.docx',
    'Brannon_Cho.docx',
    'Chih-Ta_Chen.docx',
    'Edgar_Moreau.docx',
    'Eugene_Lin.docx',
    'Jinjoo_Cho.docx',
    'Kyu_Yeon_Kim.docx',
    'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx',
    'Sirena_Huang.docx',
}
PARAGRAPH_TARGETS = {
    'Da_Zong_Guan_Master_Schedule.docx': {
        "Sep 12 dress is 09:00–12:00 (you're fine for morning)": "Sep 12 dress is 09:30–12:00 (you're fine for morning)",
    },
    'Sheraton_Front_Desk.docx': {
        'Sep 12 (Saturday): Dress rehearsal begins at 09:00 — early car at 08:30.': 'Sep 12 (Saturday): Dress rehearsal begins at 09:30.',
    },
}


def replace_paragraph(paragraph, replacements):
    before = paragraph.text
    after = before
    for old, new in replacements.items():
        after = after.replace(old, new)
    if after != before:
        paragraph.text = after
        return 1
    return 0


def update_table_time(doc):
    changes = 0
    for table in doc.tables:
        for row in table.rows:
            if not any('Dress Rehearsal' in cell.text for cell in row.cells):
                continue
            for cell in row.cells:
                if cell.text == '09:00–12:00':
                    cell.paragraphs[0].text = '09:30–12:00'
                    changes += 1
    if changes != 1:
        raise RuntimeError(f'Expected one Sep 12 dress time change, found {changes}')


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for filename in sorted(TABLE_TARGETS):
        doc = Document(SOURCE / filename)
        update_table_time(doc)
        doc.save(OUTPUT / filename)
    for filename, replacements in PARAGRAPH_TARGETS.items():
        doc = Document(SOURCE / filename)
        changes = sum(replace_paragraph(p, replacements) for p in doc.paragraphs)
        if changes != 1:
            raise RuntimeError(f'{filename}: expected one text change, found {changes}')
        doc.save(OUTPUT / filename)


if __name__ == '__main__':
    main()
