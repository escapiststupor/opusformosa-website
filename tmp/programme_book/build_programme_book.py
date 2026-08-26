from copy import deepcopy
import json
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

SOURCE = Path('/Users/pyen/Downloads/節目單/節目單文字.docx')
OUTPUT = Path('/Users/pyen/OpusFormosa/website/output/Opus音樂節_節目本_定稿版.docx')
OFFICIAL_BIOS = Path('/private/tmp/programme_book/official_full_bios.json')
EXTENDED_BIOS_A = Path('/private/tmp/programme_book/extended_bios_a.json')
EXTENDED_BIOS_B = Path('/private/tmp/programme_book/extended_bios_b.json')


def set_font(run, name='Garamond', size=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def replace_text(paragraph, text, *, font='Garamond', size=10.5, bold=False,
                 alignment=None, before=0, after=0, line=1.15):
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_font(run, font, size, bold)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def insert_after(paragraph, text='', **formatting):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    replace_text(new_para, text, **formatting)
    return new_para


def find(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f'Missing paragraph: {text}')


def add_heading_pair(doc, original, english, chinese):
    paragraph = find(doc, original)
    replace_text(paragraph, english, font='Garamond', size=13, bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=10, after=0, line=1.0)
    return insert_after(paragraph, chinese, font='Arial Unicode MS', size=13, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=5, line=1.0)


def normalize_programme(doc):
    replacements = {
        'Opening from Glassworks\tP. Glass': '“Opening” from Glassworks\tP. Glass',
        '《玻璃工程》選曲\t葛拉斯（*1937）': '〈開場〉，選自《玻璃工程》\t葛拉斯（1937–）',
        'Piano Quartet in E-flat major, Op. 47\tR. Schumann': 'Piano Quartet in E-flat Major, op. 47\tR. Schumann',
        '降E大調鋼琴四重奏，作品47\t舒曼（1810-1856）': '降 E 大調鋼琴四重奏，作品 47\t舒曼（1810–1856）',
        'Song-Book Selections & Preludes\tG. Gershwin': 'Song-Book Selections & Preludes\tG. Gershwin',
        '《歌謠集與前奏曲》選段\t蓋希文（1898-1937）': '《歌謠集》與前奏曲選段\t蓋希文（1898–1937）',
        'La Valse (two pianos)\tM. Ravel': 'La Valse (Two Pianos)\tM. Ravel',
        '《圓舞曲》（雙鋼琴）\t拉威爾（1875-1937）': '《圓舞曲》（雙鋼琴）\t拉威爾（1875–1937）',
        'Piano Concerto No. 1 in C minor, Op. 35\tD. Shostakovich': 'Concerto No. 1 for Piano, Trumpet and Strings in c minor, op. 35\tD. Shostakovich',
        'c小調第一號鋼琴協奏曲，作品35\t蕭士塔高維契（1906-1975）': 'c 小調第一號鋼琴、小號與弦樂團協奏曲，作品 35\t蕭士塔高維契（1906–1975）',
        'Much Ado About Nothing, Suite for Violin and Piano\tE. W. Korngold': 'Much Ado About Nothing, Suite for Violin and Piano\tE. W. Korngold',
        '《無事生非》（小提琴與鋼琴版組曲）\t康果爾德（1897-1957）': '《無事生非》：小提琴與鋼琴組曲\t康果爾德（1897–1957）',
        'Trio élégiaque No. 1 in G minor\tS. Rachmaninoff': 'Trio élégiaque No. 1 in g minor\tS. Rachmaninoff',
        'g小調《輓歌三重奏》第一號\t拉赫瑪尼諾夫（1873-1943）': 'g 小調第一號《輓歌三重奏》\t拉赫瑪尼諾夫（1873–1943）',
        'Cello Concerto No. 1 in C major\tJ. Haydn': 'Cello Concerto No. 1 in C Major\tJ. Haydn',
        'C大調第一號大提琴協奏曲\t海頓（1732-1809）': 'C 大調第一號大提琴協奏曲\t海頓（1732–1809）',
        'Serenade in C major for String Trio, Op. 10\tE. Dohnányi': 'Serenade in C Major for String Trio, op. 10\tE. Dohnányi',
        'C大調弦樂三重奏小夜曲，作品10\t多納尼（1877-1960）': 'C 大調弦樂三重奏小夜曲，作品 10\t多納尼（1877–1960）',
        'Piano Trio in A minor\tM. Ravel': 'Piano Trio in a minor\tM. Ravel',
        'a小調鋼琴三重奏\t拉威爾（1875-1937）': 'a 小調鋼琴三重奏\t拉威爾（1875–1937）',
        'Piano Quintet in A major, Op. 81\tA. Dvořák': 'Piano Quintet in A Major, op. 81\tA. Dvořák',
        'A大調鋼琴五重奏，作品81\t德弗札克（1841-1904）': 'A 大調鋼琴五重奏，作品 81\t德弗札克（1841–1904）',
        'Piano Quartet in A minor\tG. Mahler': 'Piano Quartet in a minor\tG. Mahler',
        'a小調鋼琴四重奏\t馬勒（1860-1911）': 'a 小調鋼琴四重奏\t馬勒（1860–1911）',
        'Piano Quartet No. 1 in G minor, Op. 15\tG. Fauré': 'Piano Quartet No. 1 in g minor, op. 15\tG. Fauré',
        'g小調第一號鋼琴四重奏，作品15\t佛瑞（1845-1924）': 'g 小調第一號鋼琴四重奏，作品 15\t佛瑞（1845–1924）',
        'String Sextet "Souvenir de Florence" in D minor, Op. 70\tP. I. Tchaikovsky': 'String Sextet “Souvenir de Florence” in d minor, op. 70\tP. I. Tchaikovsky',
        'd小調弦樂六重奏《佛羅倫斯的回憶》，作品70\t柴可夫斯基（1840-1893）': 'd 小調弦樂六重奏《佛羅倫斯的回憶》，作品 70\t柴可夫斯基（1840–1893）',
        'Duo for Violin and Viola in G Major, K. 423\tW. A. Mozart': 'Duo for Violin and Viola in G Major, K. 423\tW. A. Mozart',
        'G大調小提琴與中提琴二重奏，K.423\t莫札特（1756-1791）': 'G 大調小提琴與中提琴二重奏，K. 423\t莫札特（1756–1791）',
        'Piano Quintet in G minor\tE. Granados': 'Piano Quintet in g minor, op. 49\tE. Granados',
        'g小調鋼琴五重奏\t葛拉納多斯（1867-1916）': 'g 小調鋼琴五重奏，作品 49\t葛拉納多斯（1867–1916）',
        'String Sextet No. 1 in B-flat major, Op. 18\tJ. Brahms': 'String Sextet No. 1 in B-flat Major, op. 18\tJ. Brahms',
        '降B大調第一號弦樂六重奏，作品18\t布拉姆斯（1833-1897）': '降 B 大調第一號弦樂六重奏，作品 18\t布拉姆斯（1833–1897）',
    }
    for original, replacement in replacements.items():
        matches = [p for p in doc.paragraphs if p.text.strip() == original]
        if not matches:
            raise ValueError(f'Missing programme line: {original}')
        chinese = any('\u4e00' <= char <= '\u9fff' for char in replacement)
        for paragraph in matches:
            replace_text(paragraph, replacement, font='Arial Unicode MS' if chinese else 'Garamond', size=10.5, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.0)

    concerto_lines = [p for p in doc.paragraphs if p.text.startswith('Concerto No. 1 for Piano, Trumpet and Strings')]
    if len(concerto_lines) != 3:
        raise ValueError(f'Expected three Shostakovich concerto entries, found {len(concerto_lines)}')
    for title in concerto_lines:
        cursor = title._p.getnext()
        while cursor is not None:
            paragraph = Paragraph(cursor, title._parent)
            if paragraph.text.strip().startswith('鋼琴：'):
                insert_after(paragraph, '小號：侯傳安（Chuan-An Hou）', font='Arial Unicode MS', size=10.5,
                             bold=False, before=0, after=0, line=1.0)
                break
            cursor = cursor.getnext()


def separate_programme_pages(doc):
    concert_dates = [
        '9 月 4 日　台北　國家音樂廳',
        '9 月 9 日　高雄　衛武營國家藝術文化中心音樂廳',
        '9 月 10 日　台北　國家演奏廳',
        '9 月 12 日　台北　國家演奏廳',
        '9 月 14 日　台中　台中國家歌劇院中劇院',
        '9 月 16 日　台中　台中國家歌劇院大劇院',
    ]
    for date in concert_dates[1:]:
        find(doc, date).paragraph_format.page_break_before = True


def format_programme_notes(doc):
    pairs = [
        ('《玻璃工程》選曲／Opening from Glassworks\u3000葛拉斯', 'P. Glass: “Opening” from Glassworks', '葛拉斯：〈開場〉，選自《玻璃工程》'),
        ('降E大調鋼琴四重奏，作品47／Piano Quartet in E-flat major, Op. 47\u3000舒曼', 'R. Schumann: Piano Quartet in E-flat Major, op. 47', '舒曼：降 E 大調鋼琴四重奏，作品 47'),
        ('《歌謠集與前奏曲》選段／Song-Book Selections & Preludes\u3000蓋希文', 'G. Gershwin: Song-Book Selections & Preludes', '蓋希文：《歌謠集》與前奏曲選段'),
        ('《圓舞曲》（雙鋼琴）／La Valse (two pianos)\u3000拉威爾', 'M. Ravel: La Valse (Two Pianos)', '拉威爾：《圓舞曲》（雙鋼琴）'),
        ('c小調第一號鋼琴協奏曲，作品35／Piano Concerto No. 1 in C minor, Op. 35\u3000蕭士塔高維契', 'D. Shostakovich: Concerto No. 1 for Piano, Trumpet and Strings in c minor, op. 35', '蕭士塔高維契：c 小調第一號鋼琴、小號與弦樂團協奏曲，作品 35'),
        ('《無事生非》（小提琴與鋼琴版組曲）／Much Ado About Nothing, Suite for Violin and Piano\u3000康果爾德', 'E. W. Korngold: Much Ado About Nothing, Suite for Violin and Piano', '康果爾德：《無事生非》：小提琴與鋼琴組曲'),
        ('g小調《輓歌三重奏》第一號／Trio élégiaque No. 1 in G minor\u3000拉赫瑪尼諾夫', 'S. Rachmaninoff: Trio élégiaque No. 1 in g minor', '拉赫瑪尼諾夫：g 小調第一號《輓歌三重奏》'),
        ('C大調第一號大提琴協奏曲／Cello Concerto No. 1 in C major\u3000海頓', 'J. Haydn: Cello Concerto No. 1 in C Major', '海頓：C 大調第一號大提琴協奏曲'),
        ('C大調弦樂三重奏小夜曲，作品10／Serenade in C major for String Trio, Op. 10\u3000多納尼', 'E. Dohnányi: Serenade in C Major for String Trio, op. 10', '多納尼：C 大調弦樂三重奏小夜曲，作品 10'),
        ('a小調鋼琴三重奏／Piano Trio in A minor\u3000拉威爾', 'M. Ravel: Piano Trio in a minor', '拉威爾：a 小調鋼琴三重奏'),
        ('A大調鋼琴五重奏，作品81／Piano Quintet in A major, Op. 81\u3000德弗札克', 'A. Dvořák: Piano Quintet in A Major, op. 81', '德弗札克：A 大調鋼琴五重奏，作品 81'),
        ('a小調鋼琴四重奏／Piano Quartet in A minor\u3000馬勒', 'G. Mahler: Piano Quartet in a minor', '馬勒：a 小調鋼琴四重奏'),
        ('g小調第一號鋼琴四重奏，作品15／Piano Quartet No. 1 in G minor, Op. 15\u3000佛瑞', 'G. Fauré: Piano Quartet No. 1 in g minor, op. 15', '佛瑞：g 小調第一號鋼琴四重奏，作品 15'),
        ('d小調弦樂六重奏《佛羅倫斯的回憶》，作品70／String Sextet "Souvenir de Florence" in D minor, Op. 70\u3000柴可夫斯基', 'P. I. Tchaikovsky: String Sextet “Souvenir de Florence” in d minor, op. 70', '柴可夫斯基：d 小調弦樂六重奏《佛羅倫斯的回憶》，作品 70'),
        ('G大調小提琴與中提琴二重奏，K.423／Duo for Violin and Viola in G Major, K. 423\u3000莫札特', 'W. A. Mozart: Duo for Violin and Viola in G Major, K. 423', '莫札特：G 大調小提琴與中提琴二重奏，K. 423'),
        ('g小調鋼琴五重奏／Piano Quintet in G minor\u3000葛拉納多斯', 'E. Granados: Piano Quintet in g minor, op. 49', '葛拉納多斯：g 小調鋼琴五重奏，作品 49'),
        ('降B大調第一號弦樂六重奏，作品18／String Sextet No. 1 in B-flat major, Op. 18\u3000布拉姆斯', 'J. Brahms: String Sextet No. 1 in B-flat Major, op. 18', '布拉姆斯：降 B 大調第一號弦樂六重奏，作品 18'),
    ]
    for original, english, chinese in pairs:
        add_heading_pair(doc, original, english, chinese)


def translate_director_note(doc):
    title_pair = [p for p in doc.paragraphs if p.text.strip() in {'藝術總監的話', 'A Note from the Artistic Director'}]
    if len(title_pair) != 2:
        raise ValueError('Expected the bilingual Artistic Director title pair')
    replace_text(title_pair[0], 'A Note from the Artistic Director', font='Garamond', size=16,
                 bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0)
    replace_text(title_pair[1], '藝術總監的話', font='Arial Unicode MS', size=16,
                 bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=12, line=1.0)

    translations = {
        'Raised through a rhythm of constant departures—between Taipei, New Jersey, New York, and Philadelphia—I learned early on to find grounding in transition. Later, as my performing path unfolded across the globe, that wandering cadence carried me from intimate halls in Savannah and Little Rock to stages across Perth, Berlin, Paris, Shanghai, and Tokyo. To spend years living out of a suitcase as an only child at the piano is to know solitude intimately—a quiet world I long cherished. Yet as my journey unfolded, what ultimately endured went far beyond the private devotion of practice or the fleeting applause. It was the conversations long after the lights dimmed, the shared tables, and the strangers who, through sound, quietly became family.': '在台北、紐澤西、紐約與費城之間不斷遷徙的成長節奏裡，我很早便學會在變動中尋找安定。其後，隨著演奏生涯走向世界，這樣的流動帶我從薩凡納與小岩城的親密音樂廳，來到伯斯、柏林、巴黎、上海與東京的舞台。身為獨生子，帶著鋼琴生活在行李箱中的多年歲月，使我熟悉孤獨——那曾是我珍愛的靜謐世界。然而，旅途真正留下的，遠不只練習時私密的投入，或掌聲轉瞬即逝的光亮；而是燈光暗下後延續的談話、共同圍坐的餐桌，以及那些因聲音而悄然成為家人的陌生人。',
        'A few years ago, an unexpected family emergency called me back to Taiwan for three months—the longest stretch I had spent here since childhood. What began as an urgent visit slowly deepened into an enduring dialogue. Traveling through communities across the island, I was drawn to a rare warmth and resonance, realizing that all those nomadic years were never about drifting away, but about gathering the artistry and perspective to build something lasting right here.': '幾年前，一場突如其來的家庭緊急狀況讓我回到臺灣，停留了三個月——這是我自童年以來在此最長的一段時光。起初只是迫切的探視，漸漸卻成為一場持續的對話。走訪島上不同社群時，我被一種難得的溫暖與共鳴所吸引，並明白那些漂泊的歲月並非使我遠離；它們累積的是藝術上的養分與觀看世界的視角，讓我能在這裡建造一件長久的事。',
        'Opus Music Festival was founded on that singular impulse: to connect.': 'Opus 音樂節正是因著這個單純而堅定的念頭而成立：連結。',
        'Music possesses the rare alchemy of suspending the relentless noise of our days, aligning mind and breath within the exact same dimension of time and space. It transforms listening into a shared sanctuary. Whether you are a devoted patron whose vision helped bring this festival to life, a steadfast Friend of Opus, someone who first stumbled upon our world through a fleeting screen, or a listener seated quietly in the very back row of the hall tonight—you are an indispensable part of this living canvas.': '音樂擁有一種難得的煉金術：它能暫停日常不息的喧囂，讓心念與呼吸在同一段時間與空間裡對齊，並把聆聽化為彼此共享的庇護之所。無論您是以遠見促成音樂節誕生的支持者、一路相伴的 Opus 之友、曾在一方螢幕上偶然走進我們世界的人，或是今晚靜靜坐在音樂廳最後一排的聽眾——您都是這幅仍在生成的畫布中不可或缺的一部分。',
        'Thank you for lending your presence, your trust, and your open ears to our inaugural season. May the music we inhabit together over these days offer a rare moment of stillness, resonance, and genuine discovery.': '感謝您將您的到場、信任與敞開的耳朵，交給我們的首屆音樂節。願這幾日我們共同棲居於其中的音樂，帶來片刻難得的安靜、共鳴與真實的發現。',
    }
    for english, chinese in translations.items():
        paragraph = find(doc, english)
        insert_after(paragraph, chinese, font='Arial Unicode MS', size=10.5, bold=False,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=8, line=1.3)


def fill_bios(doc):
    records = json.loads(OFFICIAL_BIOS.read_text(encoding='utf-8'))['bios']
    proper_names = {
        '第十一屆印第安納波利斯國際小提琴大賽': '11th Indianapolis Violin Competition',
        '印第安納波利斯、蒙特婁、布宜諾斯艾利斯、Schoenfeld 與 Stulberg 等國際小提琴大賽': 'Indianapolis、Montreal、Buenos Aires、Schoenfeld 與 Stulberg 等國際小提琴大賽',
        '茱莉亞音樂院': 'Juilliard School',
        '耶魯音樂學院': 'Yale School of Music',
        '寇蒂斯音樂院': 'Curtis Institute of Music',
        '新加坡國立大學楊秀桃音樂院': 'Yong Siew Toh Conservatory of Music',
        '巴黎國立高等音樂院': 'Conservatoire de Paris',
        '維也納音樂大學': 'University of Music and Performing Arts Vienna',
        '維也納音樂暨表演藝術大學': 'University of Music and Performing Arts Vienna',
        'Detmold 音樂院': 'Hochschule für Musik Detmold',
        '首爾大學': 'Seoul National University',
        '伊莉莎白皇后': 'Queen Elisabeth',
        '克里夫蘭管弦樂團': 'Cleveland Orchestra',
    }
    def retain_proper_names(text):
        for chinese, official in proper_names.items():
            text = text.replace(chinese, official)
        return text
    by_name = {record['name']: retain_proper_names(record['bio_zh']) for record in records}
    for path in (EXTENDED_BIOS_A, EXTENDED_BIOS_B):
        by_name.update({name: retain_proper_names(bio) for name, bio in json.loads(path.read_text(encoding='utf-8')).items()})
    # Curtis's official bio does not substantiate the competition claim in the research notes.
    by_name['Steven Lin'] = by_name['Steven Lin'].replace(
        '他曾獲 Arthur Rubinstein International Piano Master Competition 獎項，並以扎實、思辨而不失即興感的詮釋建立個人風格。',
        '他以扎實、思辨而不失即興感的詮釋建立個人風格。'
    )
    by_name['Canglah Micyang'] = (
        '子晴畢業於國立陽明交通大學音樂研究所，活躍於中提琴及小提琴的演出，有時也以鋼琴合作的身份與朋友合作。她曾以獨奏家及室內樂演奏家身份參與眾多音樂節的演出，包括亞洲作曲家聯盟會議暨音樂節、台北國際現代音樂節、交大室內樂集、韓國大山音樂節、美國 Kneisel Hall 音樂節、德國 Moritzburg 音樂節、奧地利 Classics from the other side 音樂節，以及奏出福爾摩沙的年度音樂會。她同時擁有國立清華大學物理學學士及天文學碩士學位，也投入阿美族語的傳承與推廣，目前她是台灣輝達的資深工程師。'
    )
    headings = {
        '鋼琴／林易（Steven Lin）': 'Steven Lin',
        '小提琴／黃凱珉（Sirena Huang）': 'Sirena Huang',
        '中提琴／Adrien La Marca': 'Adrien La Marca',
        '大提琴／林恩俊（Eugene Lin）': 'Eugene Lin',
        '鋼琴／小林愛実（Aimi Kobayashi）': 'Aimi Kobayashi',
        '指揮／鄒佳宏（Jiahung Zou）': 'Jiahung Zou',
        '小提琴／Jinjoo Cho': 'Jinjoo Cho',
        '大提琴／Edgar Moreau': 'Edgar Moreau',
        '小提琴／丁章媛（Belle Ting）': 'Belle Ting',
        '大提琴／Brannon Cho': 'Brannon Cho',
        '鋼琴／Kyu Yeon Kim': 'Kyu Yeon Kim',
        '小提琴／Boris Borgolotto': 'Boris Borgolotto',
        '中提琴／陳志達（Chih-Ta Chen）': 'Chih-Ta Chen',
        '小提琴／王子欣（Sophie Wang）': 'Sophie Wang',
        '中提琴／嚴子晴（Canglah Micyang）': 'Canglah Micyang',
    }
    bios = {heading: by_name[name] for heading, name in headings.items()}
    for heading, bio in bios.items():
        h = find(doc, heading)
        replace_text(h, heading, font='Arial Unicode MS', size=12, bold=True, before=9, after=2, line=1.0)
        target = Paragraph(h._p.getnext(), h._parent)
        if target.text.strip() != '（簡介待補）':
            raise ValueError(f'Expected placeholder after {heading}, got {target.text!r}')
        replace_text(target, bio, font='Arial Unicode MS', size=10.5, bold=False,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8, line=1.25)

    h = find(doc, '指揮／鄒佳宏（Jiahung Zou）')
    bio = Paragraph(h._p.getnext(), h._parent)
    trumpet_heading = insert_after(bio, '小號／侯傳安（Chuan-An Hou）', font='Arial Unicode MS', size=12,
                                   bold=True, before=9, after=2, line=1.0)
    insert_after(trumpet_heading, by_name['Chuan-An Hou'],
                 font='Arial Unicode MS', size=10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 before=0, after=8, line=1.25)

    # The orchestra section is intentionally roster-only.
    orchestra_heading = find(doc, '樂團／藝響室內樂團（Opus Chamber Orchestra）')
    orchestra_bio = Paragraph(orchestra_heading._p.getnext(), orchestra_heading._parent)
    orchestra_heading._element.getparent().remove(orchestra_heading._element)
    orchestra_bio._element.getparent().remove(orchestra_bio._element)


def normalize_bilingual_text(text):
    # Chinese context uses full-width punctuation; English work titles and names retain their own punctuation.
    text = re.sub(r'\(([^()]*[\u4e00-\u9fff][^()]*)\)', r'（\1）', text)
    text = re.sub(r'\s*（\s*', '（', text)
    text = re.sub(r'\s*）', '）', text)
    text = re.sub(r'(?<=[\u4e00-\u9fff])\s*,\s*', '，', text)
    text = re.sub(r',\s*(?=[\u4e00-\u9fff])', '，', text)
    text = re.sub(r'(?<=[\u4e00-\u9fff])\s*:\s*', '：', text)
    text = re.sub(r'\s*，\s*', '，', text)
    text = re.sub(r'\s*。\s*', '。', text)
    text = re.sub(r'\s*；\s*', '；', text)
    text = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])', '', text)
    # Insert one half-width space at direct CJK/Latin boundaries. Punctuation and brackets remain tight.
    text = re.sub(r'(?<=[\u4e00-\u9fff])(?=[A-Za-z])', ' ', text)
    text = re.sub(r'(?<=[A-Za-z])(?=[\u4e00-\u9fff])', ' ', text)
    return text


def normalize_bilingual_punctuation(doc):
    changed = 0
    for paragraph in doc.paragraphs:
        original = paragraph.text
        normalized = normalize_bilingual_text(original)
        if normalized == original:
            continue
        nonempty_runs = [run for run in paragraph.runs if run.text]
        if len(nonempty_runs) == 1:
            nonempty_runs[0].text = normalized
        else:
            # These paragraphs have no intentional inline style transitions in the source.
            # Preserve the first run's formatting while replacing only its text payload.
            if nonempty_runs:
                nonempty_runs[0].text = normalized
                for run in nonempty_runs[1:]:
                    run.text = ''
            else:
                paragraph.add_run(normalized)
        changed += 1
    return changed


def localize_proper_names(doc):
    replacements = {
        'Artist Diploma': '藝術家文憑',
        'Performance Diploma': '演奏文憑',
        'The Juilliard School': '茱莉亞音樂院',
        'Juilliard School': '茱莉亞音樂院',
        'Curtis Institute of Music': '寇蒂斯音樂院',
        'Conservatoire National Supérieur de Paris': '巴黎高等音樂學院',
        'University of Music and Performing Arts Vienna': '維也納表演藝術大學',
        'Yong Siew Toh Conservatory of Music': '新加坡國立大學楊秀桃音樂院',
        'New England Conservatory': '新英格蘭音樂院',
        'Cleveland Institute of Music': '克里夫蘭音樂院',
        'Manhattan School of Music': '曼哈頓音樂院',
        'Seoul National University': '首爾大學',
        'Northwestern University Bienen School of Music': '西北大學Bienen音樂院',
        'International Arthur Grumiaux Competition': '葛羅米歐小提琴大賽大獎',
        'International Cooper Violin Competition': '國際 Cooper 小提琴大賽',
        'International Brahms Competition': '國際布拉姆斯小提琴大賽',
        'International Isaac Stern Competition': '國際 Isaac Stern 小提琴大賽',
        'International Khachaturian Competition': '國際哈察圖良小提琴大賽',
        'Asia-Pacific International Fryderyk Chopin Piano Competition': '亞太國際蕭邦鋼琴大賽',
        'Asian International Fryderyk Chopin Piano Competition': '亞洲國際蕭邦鋼琴大賽',
        'International Chopin Piano Competition': '國際蕭邦鋼琴大賽',
        'Indianapolis International Violin Competition': '印第安納波利斯國際小提琴大賽',
        'Montreal International Musical Competition': '蒙特婁國際音樂大賽',
        'Buenos Aires International Violin Competition': '布宜諾斯艾利斯國際小提琴大賽',
        'Schoenfeld International String Competition': 'Schoenfeld 國際弦樂大賽',
        'Stulberg International String Competition': 'Stulberg 國際弦樂大賽',
        'Hiroshima International Conducting Competition': '廣島國際指揮大賽',
        'Jeju International Brass Competition': '濟州國際銅管大賽',
        'William Primrose Competition': 'William Primrose 中提琴大賽',
        'Lionel Tertis International Viola Competition': 'Lionel Tertis 國際中提琴大賽',
        'International Johannes Brahms Competition': '國際布拉姆斯大賽',
        'International Tchaikovsky Competition for Young Musicians': '國際柴科夫斯基青年音樂家大賽',
        'Johansen International Competition': 'Johansen 國際大賽',
        'Geneva International Music Competition': '日內瓦國際音樂大賽',
        'ARD International Music Competition': 'ARD 國際音樂大賽',
        'Rostropovich Cello Competition': '羅斯托波維奇大提琴大賽',
        'International Tchaikovsky Competition': '國際柴科夫斯基大賽',
        'International Paulo Cello Competition': 'Paulo 國際大提琴大賽',
        'Fischoff National Chamber Music Competition': 'Fischoff 全國室內樂大賽',
        'Melbourne International Chamber Music Competition': '墨爾本國際室內樂大賽',
        'Osaka International Chamber Music Competition': '大阪國際室內樂大賽',
        'International Wolfgang Marschner Competition': '國際 Wolfgang Marschner 大賽',
        'Paul Hindemith Competition Berlin': 'Paul Hindemith 柏林大賽',
        'Louis Spohr Competition Weimar': 'Louis Spohr 威瑪大賽',
        'Andrea Postacchini Violin Competition': 'Andrea Postacchini 小提琴大賽',
        'Taipei International New Music Festival': '台北國際現代音樂節',
        'National Trumpet Competition': '全美小號大賽',
        'Young Concert Artists International Auditions': 'Young Concert Artists 國際甄選',
        'International Violin Competition': '國際小提琴大賽',
        'Violin Competition': '小提琴大賽',
        'Dublin International Piano Competition': 'Dublin 國際鋼琴大賽',
        'Queen Elisabeth International Music Competition': '伊莉莎白皇后國際音樂大賽',
        'Cleveland International Piano Competition': 'Cleveland 國際鋼琴大賽',
        'Gina Bachauer International Young Artists Piano Competition': 'Gina Bachauer 國際青年鋼琴大賽',
        'Geneva International Music Competition': '日內瓦國際音樂大賽',
        'Missouri Southern International Piano Competition': '密蘇里南方國際鋼琴大賽',
        'Konzertexamen': '最高演奏文憑',
        'Cleveland Symphony Orchestra': '克里夫蘭交響樂團',
        'National Taiwan Symphony Orchestra': '國立臺灣交響樂團',
        'Shanghai Symphony Orchestra': '上海交響樂團',
        'Odense Symphony Orchestra': '歐登塞交響樂團',
        'New Zealand Symphony Orchestra': '紐西蘭交響樂團',
    }
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
        text = paragraph.text
        for english, chinese in replacements.items():
            text = text.replace(english, chinese)
        if text != paragraph.text:
            replace_text(paragraph, text, font='Arial Unicode MS', size=10.5, bold=False,
                         alignment=paragraph.alignment, before=paragraph.paragraph_format.space_before.pt or 0,
                         after=paragraph.paragraph_format.space_after.pt or 0,
                         line=paragraph.paragraph_format.line_spacing or 1.15)


def main():
    doc = Document(SOURCE)
    translate_director_note(doc)
    normalize_programme(doc)
    separate_programme_pages(doc)
    format_programme_notes(doc)
    fill_bios(doc)
    normalize_bilingual_punctuation(doc)
    localize_proper_names(doc)
    normalize_bilingual_punctuation(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == '__main__':
    main()
