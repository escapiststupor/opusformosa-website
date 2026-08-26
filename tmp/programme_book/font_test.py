from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

doc = Document()
run = doc.add_paragraph().add_run('未指定字型：測試中文顯示，藝術總監的話。')
run.font.size = Pt(14)
run._element.get_or_add_rPr().set(qn('w:lang'), 'zh-TW')
for font in ['Arial Unicode MS', 'PingFang TC', 'STFangsong', 'LiSong Pro', 'Songti TC']:
    run = doc.add_paragraph().add_run(font + '：測試中文顯示，藝術總監的話。')
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), font)
    run._element.get_or_add_rPr().set(qn('w:lang'), 'zh-TW')
    run.font.size = Pt(14)
doc.save('tmp/programme_book/font-test.docx')
