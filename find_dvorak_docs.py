from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')


def text_of(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return '\n'.join(parts)


for path in sorted(ROOT.glob('*.docx')):
    text = text_of(Document(path))
    if 'Dvořák' in text or 'Dvorak' in text:
        print(path.name)
        for line in text.splitlines():
            if 'Dvořák' in line or 'Dvorak' in line:
                print('  ' + line)
