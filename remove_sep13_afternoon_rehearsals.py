from pathlib import Path
from docx import Document


LOGISTICS = Path("/Users/pyen/OpusFormosa/festival_planning/logistics")


def remove_row(table, index):
    table._tbl.remove(table.rows[index]._tr)


def cell_text(row):
    return [cell.text.replace("\n", " ").strip() for cell in row.cells]


def remove_sep13_afternoon_rows(document_path):
    document = Document(document_path)
    removed = []
    for table_number, table in enumerate(document.tables):
        active_date = ""
        targets = []
        for row_number, row in enumerate(table.rows):
            cells = cell_text(row)
            if cells and cells[0] in {"Sep 13", "Sep13", "September 13"}:
                active_date = "Sep 13"
            elif cells and cells[0].startswith("Sep "):
                active_date = cells[0]

            event_text = " ".join(cells)
            if (
                active_date == "Sep 13"
                and len(cells) >= 3
                and cells[1] in {"13:00–15:00", "15:00–17:00"}
                and "Brahms: String Sextet No.1 Op.18" in event_text
            ):
                targets.append(row_number)

        for row_number in reversed(targets):
            removed.append((table_number, row_number, cell_text(table.rows[row_number])))
            remove_row(table, row_number)

    document.save(document_path)
    return removed


for filename in ("Sirena_Huang.docx", "Eugene_Lin.docx"):
    path = LOGISTICS / filename
    changed = remove_sep13_afternoon_rows(path)
    if len(changed) not in {0, 2}:
        raise RuntimeError(f"Expected zero or two Sep 13 afternoon rows in {filename}; found {changed!r}")
    document = Document(path)
    labeled = []
    for table_number, table in enumerate(document.tables):
        for row_number, row in enumerate(table.rows):
            cells = cell_text(row)
            if (
                len(cells) >= 3
                and cells[0] == ""
                and cells[1] == "19:30–21:30"
                and "Brahms: String Sextet No.1 Op.18" in " ".join(cells)
            ):
                row.cells[0].text = "Sep 13"
                labeled.append((table_number, row_number))
    if len(labeled) not in {0, 1}:
        raise RuntimeError(f"Expected zero or one remaining Sep 13 evening row in {filename}; found {labeled!r}")
    document.save(path)
    print(filename, "removed", changed, "labeled", labeled)

# The master handbook's daily schedule uses time (rather than a repeated date)
# in its first column. Delete its one Sep 13 entry for Eugene and Sirena.
master = LOGISTICS / "Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"
document = Document(master)
changed = []
for table_number, table in enumerate(document.tables):
    targets = []
    for row_number, row in enumerate(table.rows):
        cells = cell_text(row)
        if (
            len(cells) >= 3
            and cells[0] == "13:00–15:00"
            and "Eugene Lin; Sirena Huang" in " ".join(cells)
            and "Brahms: String Sextet No.1 Op.18" in " ".join(cells)
        ):
            targets.append(row_number)
    for row_number in reversed(targets):
        changed.append((table_number, row_number, cell_text(table.rows[row_number])))
        remove_row(table, row_number)
if len(changed) not in {0, 1}:
    raise RuntimeError(f"Expected zero or one Sep 13 afternoon block in master handbook; found {changed!r}")
document.save(master)
print(master.name, "removed", changed)
