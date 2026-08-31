from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path("/Users/pyen/OpusFormosa/festival_planning/logistics")


def compact_font(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "PingFang TC"
                run.font.size = Pt(8)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find("w:cantSplit", tr_pr.nsmap) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


# Edgar's individual itinerary: make the airport-to-hotel plan visible before his first rehearsal.
path = ROOT / "Edgar_Moreau.docx"
document = Document(path)
table = document.tables[1]
arrival_row = next(
    (row for row in table.rows if "half-day early check-in arranged" in row.cells[2].text),
    None,
)
if arrival_row is None:
    first_row = table.rows[1]
    copied_tr = deepcopy(first_row._tr)
    position = list(table._tbl).index(first_row._tr)
    table._tbl.insert(position, copied_tr)
    arrival_row = next(row for row in table.rows if row._tr is copied_tr)
    for cell, value in zip(arrival_row.cells, [
        "Sep 8",
        "06:50–~08:30",
        "Airport arrival — BR88; half-day early check-in arranged",
        "Taipei Taoyuan Airport → Sheraton Grand Taipei",
    ]):
        cell.text = value
    compact_font(arrival_row)
keep_row_together(arrival_row)
document.save(path)

# Sheraton: hotel-facing instruction in Edgar's arrival cell.
path = ROOT / "Sheraton_Front_Desk.docx"
document = Document(path)
for table in document.tables:
    for row in table.rows:
        if row.cells and row.cells[0].text == "Edgar Moreau":
            arrival = row.cells[1]
            if "Half-day early check-in" not in arrival.text:
                paragraph = arrival.add_paragraph()
                run = paragraph.add_run("Half-day early check-in arranged upon arrival.")
                run.font.name = "PingFang TC"
                run.font.size = Pt(8)
document.save(path)

# Main handbook: both airport-control and daily-run-sheet entries carry the instruction.
path = ROOT / "Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"
document = Document(path)
for table_index in (2, 10):
    for row in document.tables[table_index].rows[1:]:
        if "Edgar Moreau" in row.cells[1].text and "BR88" in " | ".join(cell.text for cell in row.cells):
            row.cells[-1].text = "接機；安排半日早入住"
            compact_font(row)
document.save(path)

print("Added Edgar's half-day early check-in instructions.")
