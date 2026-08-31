from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')


personal_path = ROOT / 'Kyu_Yeon_Kim.docx'
personal = Document(personal_path)
table = next(t for t in personal.tables if [c.text.strip() for c in t.rows[0].cells][:4] == ['Date', 'Time', 'Programme / Event', 'Venue'])
obsolete = [row for row in table.rows if row.cells[1].text.strip() == '16:00–19:00' and 'Practice room' in row.cells[2].text]
if len(obsolete) != 1:
    raise RuntimeError(f'Expected one obsolete personal-schedule row, found {len(obsolete)}')
table._tbl.remove(obsolete[0]._tr)
personal.save(personal_path)

handbook_path = ROOT / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx'
handbook = Document(handbook_path)
replacements = 0
for table in handbook.tables:
    for row in table.rows:
        values = [cell.text for cell in row.cells]
        if 'Kyu Yeon Kim' not in values:
            continue
        for cell in row.cells:
            if '16:00-19:00' in cell.text:
                cell.text = cell.text.replace('16:00-19:00', '14:00-19:00')
                replacements += 1
if replacements != 1:
    raise RuntimeError(f'Expected one remaining handbook Kyu practice row, found {replacements}')
handbook.save(handbook_path)
print('Removed obsolete Kyu personal row and corrected handbook duplicate timeline.')
