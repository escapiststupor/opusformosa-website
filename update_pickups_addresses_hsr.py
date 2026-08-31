from pathlib import Path

from docx import Document
from docx.shared import Pt


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/pickups_addresses_hsr_update')
ADDRESS_TARGETS = {
    'Adrien_La_Marca.docx', 'Aimi_Kobayashi.docx', 'Boris_Borgolotto.docx',
    'Brannon_Cho.docx', 'Da_Zong_Guan_Master_Schedule.docx', 'Edgar_Moreau.docx',
    'Eugene_Lin.docx', 'Jinjoo_Cho.docx', 'Kyu_Yeon_Kim.docx',
    'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx', 'Sheraton_Front_Desk.docx',
    'Sirena_Huang.docx', 'Steven_Lin.docx',
}
FOREIGN_ARTISTS = {
    'Aimi_Kobayashi.docx', 'Adrien_La_Marca.docx', 'Boris_Borgolotto.docx',
    'Brannon_Cho.docx', 'Edgar_Moreau.docx', 'Jinjoo_Cho.docx', 'Kyu_Yeon_Kim.docx',
}
HSR_GROUPS = {
    'Edgar_Moreau.docx': [
        'Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.',
        'Sep 13, Taichung: Edgar Moreau and Steven Lin.',
        'Sep 14, Taichung: Edgar Moreau, Boris Borgolotto, and Steven Lin.',
        'Sep 16, Taichung: Edgar Moreau and Steven Lin.',
    ],
    'Jinjoo_Cho.docx': ['Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.'],
    'Boris_Borgolotto.docx': ['Sep 14, Taichung: Boris Borgolotto, Edgar Moreau, and Steven Lin.'],
    'Steven_Lin.docx': [
        'Sep 13, Taichung: Steven Lin and Edgar Moreau.',
        'Sep 14, Taichung: Steven Lin, Edgar Moreau, and Boris Borgolotto.',
        'Sep 16, Taichung: Steven Lin and Edgar Moreau.',
    ],
    'Da_Zong_Guan_Master_Schedule.docx': [
        'Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.',
        'Sep 13, Taichung: Edgar Moreau and Steven Lin.',
        'Sep 14, Taichung: Edgar Moreau, Boris Borgolotto, and Steven Lin.',
        'Sep 16, Taichung: Edgar Moreau and Steven Lin.',
    ],
    'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx': [
        'Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.',
        'Sep 13, Taichung: Edgar Moreau and Steven Lin.',
        'Sep 14, Taichung: Edgar Moreau, Boris Borgolotto, and Steven Lin.',
        'Sep 16, Taichung: Edgar Moreau and Steven Lin.',
    ],
    'Sheraton_Front_Desk.docx': [
        'Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.',
        'Sep 13, Taichung: Edgar Moreau and Steven Lin.',
        'Sep 14, Taichung: Edgar Moreau, Boris Borgolotto, and Steven Lin.',
        'Sep 16, Taichung: Edgar Moreau and Steven Lin.',
    ],
}

CHR_ADDRESS = 'National Concert Hall Backstage, Basement Floor, No. 21-1, Zhongshan S. Rd., Zhongzheng Dist., Taipei City 100012, Taiwan'
SHERATON = 'Sheraton Grand Taipei — No. 12, Sec. 1, Zhongxiao East Road, Zhongzheng District, Taipei, Taiwan 10049 — +886-2-2321-5511'
PICKUP = ('Arrival pickup: A local team member will be waiting in the arrivals area with a sign showing your name. '
          'If you cannot find the representative, contact Phyllis Canglah M. or Ethan (Yiting) via WhatsApp. ')


def all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def set_text(paragraph, text):
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)


def update_document(source, output):
    doc = Document(source)
    filename = source.name
    pickup_added = False
    for paragraph in all_paragraphs(doc):
        text = paragraph.text
        if text.startswith('CHR rehearsal rooms'):
            set_text(paragraph, 'All CHR rehearsal rooms (CHR1 / CHR2 / CHR3 / CHR4) — ' + CHR_ADDRESS + '. Use the staff entrance: https://maps.app.goo.gl/kxPsa1KHaKpFppgi9. Please bring your ID card or passport and show it to the guard before entering.')
        elif text.startswith('CHR1–4:'):
            set_text(paragraph, 'CHR1–4: ' + CHR_ADDRESS)
        elif text.startswith('Sheraton:'):
            set_text(paragraph, 'Sheraton: ' + SHERATON)
        elif filename in FOREIGN_ARTISTS and text.startswith('Note:') and not pickup_added:
            set_text(paragraph, 'Note: ' + PICKUP + text.removeprefix('Note: ').lstrip())
            pickup_added = True
    if filename in FOREIGN_ARTISTS and not pickup_added:
        raise RuntimeError(f'{filename}: no Note paragraph found for arrival-pickup instruction')
    if filename in HSR_GROUPS:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run('HSR travel companions')
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        for line in HSR_GROUPS[filename]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(line)
            run.font.name = 'Arial'
            run.font.size = Pt(9)
    doc.save(output)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for filename in sorted(ADDRESS_TARGETS):
        update_document(SOURCE / filename, OUTPUT / filename)


if __name__ == '__main__':
    main()
