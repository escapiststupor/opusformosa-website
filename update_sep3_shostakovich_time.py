from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/sep3_shostakovich_update')
PERSONAL_FILES = [
    'Chih-Ta_Chen.docx',
    'Eugene_Lin.docx',
    'Hou_Chuan-An_Trumpet.docx',
    'Sirena_Huang.docx',
    'Steven_Lin.docx',
]


def update_personal(path):
    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 3 and 'Shostakovich: Piano Concerto No.1' in row.cells[2].text and row.cells[1].text.strip() == '19:30–21:30':
                row.cells[1].text = '19:15–21:45'
    return doc


def update_handbook():
    doc = Document(SOURCE / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')
    roster = 'Chih-Ta Chen; Eugene Lin; Hou Chuan-An (Trumpet); Sirena Huang; Steven Lin'
    for table in doc.tables:
        for row in table.rows:
            values = ' | '.join(cell.text for cell in row.cells)
            if 'Shostakovich: Piano Concerto No.1' in values and '台北室內合唱團大排練室' in values:
                row.cells[0].text = '19:15–21:45'
                row.cells[1].text = roster
    return doc


def update_master_schedule():
    doc = Document(SOURCE / 'Da_Zong_Guan_Master_Schedule.docx')
    for table in doc.tables:
        next_sep4 = next((row for row in table.rows if row.cells[0].text.strip() == 'Sep 4'), None)
        sep3_entry = next((row for row in table.rows if row.cells[1].text.strip() == '14:30' and 'Steinway center' in row.cells[2].text), None)
        if next_sep4 is None or sep3_entry is None:
            continue
        new_tr = deepcopy(sep3_entry._tr)
        table._tbl.insert(table._tbl.index(next_sep4._tr), new_tr)
        new_row = next(row for row in table.rows if row._tr is new_tr)
        new_row.cells[0].text = ''
        new_row.cells[1].text = '19:15–21:45'
        new_row.cells[2].text = 'Shostakovich rehearsal — Chih-Ta Chen, Eugene Lin, Hou Chuan-An (Trumpet), Sirena Huang, and Steven Lin. 台北室內合唱團大排練室.'
        break
    return doc


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for name in PERSONAL_FILES:
        update_personal(SOURCE / name).save(OUTPUT / name)
    update_handbook().save(OUTPUT / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')
    update_master_schedule().save(OUTPUT / 'Da_Zong_Guan_Master_Schedule.docx')


if __name__ == '__main__':
    main()
