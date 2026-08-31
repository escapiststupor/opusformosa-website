from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')


def text_of(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return '\n'.join(parts)


for path in sorted(ROOT.glob('*.docx')):
    lines = [line for line in text_of(Document(path)).splitlines()
             if any(word in line.lower() for word in ('sheraton', 'taxi', 'car'))]
    if lines:
        print(f'---{path.name}---')
        print('\n'.join(lines))
