from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PATH = "/Users/pyen/OpusFormosa/festival_planning/logistics/Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"
FONT = "PingFang TC"


def set_run_font(run):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)
    rfonts.set(qn("w:hint"), "eastAsia")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:eastAsia"), "zh-TW")


def visit_paragraphs(container):
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                visit_paragraphs(cell)


document = Document(PATH)
visit_paragraphs(document)
for section in document.sections:
    visit_paragraphs(section.header)
    visit_paragraphs(section.footer)

for style in document.styles:
    if style.type in {1, 2, 3}:
        style.font.name = FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr._add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), FONT)
        rfonts.set(qn("w:hint"), "eastAsia")

document.save(PATH)
print("Applied PingFang TC and Traditional Chinese language metadata to handbook text and styles.")
