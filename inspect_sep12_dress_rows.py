from pathlib import Path

from docx import Document


ROOT = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')


for path in sorted(ROOT.glob('*.docx')):
    doc = Document(path)
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' | ') for cell in row.cells]
            row_text = ' '.join(cells).lower()
            if 'sep 12' in row_text and ('dress rehearsal' in row_text or 'rehearsal' in row_text):
                print(path.name, table_index, row_index, ' || '.join(cells))
            elif '09:00' in row_text and 'dress rehearsal' in row_text:
                print(path.name, table_index, row_index, 'CHECK DATE CONTEXT || ' + ' || '.join(cells))
    paragraphs = '\n'.join(p.text for p in doc.paragraphs)
    for line in paragraphs.splitlines():
        if 'Sep 12' in line and '09:00' in line:
            print(path.name, 'PARAGRAPH || ' + line)
