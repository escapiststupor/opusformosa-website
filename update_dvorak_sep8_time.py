from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/dvorak_sep8_update')
TARGETS = {
    'Adrien_La_Marca.docx',
    'Belle_Ting.docx',
    'Boris_Borgolotto.docx',
    'Edgar_Moreau.docx',
    'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx',
    'Steven_Lin.docx',
}


def update_document(source, output):
    doc = Document(source)
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            if not any('Dvořák: Piano Quintet Op.81' in cell.text for cell in row.cells):
                continue
            for cell in row.cells:
                if cell.text == '19:30–21:30':
                    cell.paragraphs[0].text = '19:00–21:30'
                    changed += 1
    if changed != 1:
        raise RuntimeError(f'{source.name}: expected one Dvořák time change, found {changed}')
    doc.save(output)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for filename in sorted(TARGETS):
        update_document(SOURCE / filename, OUTPUT / filename)


if __name__ == '__main__':
    main()
