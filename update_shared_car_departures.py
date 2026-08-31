from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/shared_car_departures_update')


def update_kyu():
    doc = Document(SOURCE / 'Kyu_Yeon_Kim.docx')
    for p in doc.paragraphs:
        if 'shared 12:30 taxi' in p.text:
            p.text = p.text.replace('shared 12:30 taxi', 'shared 12:40 taxi')
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 3 and row.cells[1].text.strip() == '12:30' and 'Shared taxi with Adrien' in row.cells[2].text:
                row.cells[1].text = '12:40'
    return doc


def update_front_desk():
    doc = Document(SOURCE / 'Sheraton_Front_Desk.docx')
    cars = doc.tables[1]
    for row in cars.rows:
        if len(row.cells) < 5 or row.cells[1].text.strip() != '12:30':
            continue
        destination = row.cells[3].text
        notes = row.cells[4].text
        if ('CHR' in destination or 'Taipei Concert Hall' in destination or 'Taipei Recital Hall' in destination) and ('Dress Rehearsal' in notes or 'Fauré' in notes or 'Tchaikovsky' in notes or 'First day' in notes or 'priority check-in' in notes or notes == ''):
            row.cells[1].text = '12:40'
    return doc


def update_master():
    doc = Document(SOURCE / 'Da_Zong_Guan_Master_Schedule.docx')
    for table in doc.tables:
        for row in table.rows:
            text = ' | '.join(cell.text for cell in row.cells)
            if 'shared 12:30 car' in text:
                for cell in row.cells:
                    if 'shared 12:30 car' in cell.text:
                        cell.text = cell.text.replace('shared 12:30 car', 'shared 12:40 car')
            if len(row.cells) >= 3 and row.cells[1].text.strip() == '12:30' and ('→ CHR' in row.cells[2].text or 'Taipei Concert Hall' in row.cells[2].text or 'Taipei Recital Hall' in row.cells[2].text):
                row.cells[1].text = '12:40'
    return doc


def update_handbook():
    doc = Document(SOURCE / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'shared 12:30 car' in cell.text:
                    cell.text = cell.text.replace('shared 12:30 car', 'shared 12:40 car')
    return doc


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    update_kyu().save(OUTPUT / 'Kyu_Yeon_Kim.docx')
    update_front_desk().save(OUTPUT / 'Sheraton_Front_Desk.docx')
    update_master().save(OUTPUT / 'Da_Zong_Guan_Master_Schedule.docx')
    update_handbook().save(OUTPUT / 'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx')


if __name__ == '__main__':
    main()
