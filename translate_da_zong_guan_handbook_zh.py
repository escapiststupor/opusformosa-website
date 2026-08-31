import re

from docx import Document


PATH = "/Users/pyen/OpusFormosa/festival_planning/logistics/Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"


EXACT = {
    "Da Zong Guan (Chief Coordinator) - Master Operations Handbook": "大總管－總務運作手冊",
    "Who is where, when, and doing what": "掌握每位人員的所在、時間與工作事項",
    "Festival dates: 2-17 September 2026 | Version: operational working copy": "音樂節日期：2026年9月2日至17日｜版本：現場執行工作版",
    "Control Desk Essentials": "大總管工作重點",
    "Foreign Artist Travel and Hotel Control": "外籍音樂家交通與住宿管理",
    "Fixed Special Commitments": "固定特別事項",
    "Daily Run Sheet": "每日執行表",
    "Read each day from top to bottom. Airport movements and special commitments are included alongside rehearsal, dress rehearsal, concert, and travel commitments.": "請依每日表格由上而下執行；表內包含機場接送、特別事項、排練、彩排、演出與交通安排。",
    "People Index": "人員索引",
    "Venue Reference": "場館參考",
    "Key locations & access": "重要地點與出入方式",
    "HSR travel companions": "高鐵同行人員",
    "Item": "項目",
    "Instruction": "說明",
    "When": "時間",
    "Artist": "音樂家",
    "Hotel": "住宿飯店",
    "Flight / Terminal": "航班／航廈",
    "Action": "工作事項",
    "Who": "人員",
    "What": "事項",
    "Where": "地點",
    "Control note": "執行備註",
    "Time": "時間",
    "Person": "人員",
    "Activity": "活動",
    "Control action": "執行事項",
    "Hotel / Base": "住宿／所在地",
    "Concert dates": "演出日期",
    "Venue": "場館",
    "Operational use / access": "用途／出入方式",
    "Local contacts": "當地聯絡人",
    "Sheraton guests": "喜來登住宿音樂家",
    "Boris": "Boris",
    "NTCH rooms": "國家音樂廳排練室",
    "Kaohsiung": "高雄",
    "Kaohsiung poster task": "高雄海報任務",
    "Southbound travel": "南下交通",
    "Airport pickup": "接機",
    "Airport drop-off": "送機",
    "Sponsor Dinner (must arrive)": "贊助人晚宴（務必出席）",
    "Sponsor Dinner (may arrive later)": "贊助人晚宴（可較晚抵達）",
    "Confirm attendance": "確認出席",
    "Personal practice": "個人練習",
    "Only upright-piano session": "唯一使用直立鋼琴的時段",
    "HSR": "高鐵",
    "HSR day trip": "高鐵當日往返",
    "Local / not listed": "本地／未列出",
    "Not listed": "未列出",
    "Local (Taipei)": "本地（台北）",
    "Hotel in Taipei (own arrangement)": "台北住宿（自行安排）",
    "TBD": "待確認",
}


REPLACEMENTS = [
    ("All CHR rehearsal rooms (CHR1 / CHR2 / CHR3 / CHR4) — National Concert Hall Backstage, Basement Floor, No. 21-1, Zhongshan S. Rd., Zhongzheng Dist., Taipei City 100012, Taiwan. Use the staff entrance: https://maps.app.goo.gl/kxPsa1KHaKpFppgi9. Please bring your ID card or passport and show it to the guard before entering.", "所有 CHR 排練室（CHR1／CHR2／CHR3／CHR4）均位於國家音樂廳後台地下樓層：台北市中正區中山南路21-1號。請由工作人員入口進入：https://maps.app.goo.gl/kxPsa1KHaKpFppgi9；入場前請攜帶身分證或護照並出示給警衛。"),
    ("National Concert Hall / National Recital Hall — No. 21-1, Zhongshan S. Rd., Zhongzheng Dist., Taipei City 100012, Taiwan.", "國家音樂廳／國家演奏廳：台北市中正區中山南路21-1號。"),
    ("National Taichung Theater — 101, Huilai Rd., Sec. 2, Xitun District, Taichung City 407025, Taiwan.", "臺中國家歌劇院：台中市西屯區惠來路二段101號。"),
    ("National Kaohsiung Center for the Arts (Weiwuying) — No. 1, Sanduo 1st Rd., Fengshan Dist., Kaohsiung City 830043, Taiwan.", "衛武營國家藝術文化中心：高雄市鳳山區三多一路1號。"),
    ("Sponsor's Dinner — 1F., No. 70-1, Chengde Rd., Sec. 1, Datong Dist., Taipei City, Taiwan.", "贊助人晚宴：台北市大同區承德路一段70-1號1樓。"),
    ("Sep 9, Kaohsiung: Jinjoo Cho and Edgar Moreau.", "9月9日，高雄：Jinjoo Cho、Edgar Moreau。"),
    ("Sep 13, Taichung: Edgar Moreau and Steven Lin.", "9月13日，台中：Edgar Moreau、Steven Lin。"),
    ("Sep 14, Taichung: Edgar Moreau, Boris Borgolotto, and Steven Lin.", "9月14日，台中：Edgar Moreau、Boris Borgolotto、Steven Lin。"),
    ("Sep 16, Taichung: Edgar Moreau and Steven Lin.", "9月16日，台中：Edgar Moreau、Steven Lin。"),
    ("CHR1-CHR4 = National Concert Hall backstage rehearsal rooms. Rental blocks: morning 09:00-12:00; afternoon 13:00-17:00; evening 18:00-22:00.", "CHR1-CHR4 為國家音樂廳後台排練室。租用時段：上午09:00-12:00；下午13:00-17:00；晚上18:00-22:00。"),
    ("Weiwuying performer entrance: scan credential barcode; stage manager Hao-Xian distributes backstage passes.", "衛武營演職人員入口：刷證件條碼入場；舞台監督豪賢發放後台證。"),
    ("Da Zong Guan: bring the festival posters on the southbound trip to Kaohsiung.", "大總管：南下高雄時請攜帶音樂節海報。"),
    ("Kaohsiung and Taichung are HSR day trips. Confirm train number, carriage/seat, and station meeting point before departure; do not assume tickets are booked.", "高雄與台中均為高鐵當日往返。出發前請確認車次、車廂／座位及車站集合點；不可假設車票已訂妥。"),
    ("Leave stage passes at the front desk; remind each guest to bring the pass to rehearsal.", "請將後台證留在櫃檯，並提醒每位住客攜帶後台證前往排練。"),
    ("Goodmore Hotel. Arrival pickup required.", "谷墨商旅。需安排接機。"),
    ("Practice room: grand piano", "琴房：平台鋼琴"),
    ("Practice room: upright piano", "琴房：直立鋼琴"),
    ("Dress Rehearsal (chamber)", "彩排（室內樂）"),
    ("Dress Rehearsal (orchestra)", "彩排（管弦樂）"),
    ("Dress Rehearsal", "彩排"),
    ("Orchestra Rehearsal (Haydn)", "管弦樂排練（Haydn）"),
    ("CONCERT - Opening Night Concert", "音樂會－開幕之夜"),
    ("CONCERT - Across Generations: Trio to Concerto", "音樂會－跨越世代：從三重奏到協奏曲"),
    ("CONCERT - Chamber Series I", "音樂會－室內樂系列 I"),
    ("CONCERT - Chamber Series II", "音樂會－室內樂系列 II"),
    ("CONCERT - Chamber Series III", "音樂會－室內樂系列 III"),
    ("CONCERT - Closing Night Concert", "音樂會－閉幕之夜"),
    ("Private event in Taichung", "台中私人活動"),
    ("HSR to Kaohsiung", "高鐵前往高雄"),
    ("HSR to Taichung", "高鐵前往台中"),
    ("HSR return", "高鐵返回"),
    ("Confirm call time, passes, and instrument needs", "確認集合時間、後台證與樂器需求"),
    ("Confirm ticket and station handoff", "確認車票與車站交接"),
    ("Airport pickup; priority Sheraton check-in and stage pass; shared 12:40 car with Adrien to CHR4", "接機；優先辦理喜來登入住與領取後台證；12:40與 Adrien 共乘前往 CHR4。"),
    ("Main Taipei concert venue. Stage passes required for rehearsal space access.", "台北主要演出場館；進入排練空間需持後台證。"),
    ("NTCH backstage rehearsal rooms.", "國家音樂廳後台排練室。"),
    ("Dress rehearsals and Chamber Series I & II concerts.", "彩排及室內樂系列 I、II 演出場地。"),
    ("Kaohsiung dress rehearsals and Sep 9 concert; performer entrance and barcode credential.", "高雄彩排與9月9日演出場地；由演職人員入口刷證件條碼進入。"),
    ("Sep 14 dress rehearsal and Chamber Series III concert.", "9月14日彩排與室內樂系列 III 演出場地。"),
    ("Sep 16 closing concert dress rehearsal and concert.", "9月16日閉幕音樂會彩排與演出場地。"),
    ("Sep 2 sponsor dinner. Core arrival 17:00; Aimi may arrive 18:00.", "9月2日贊助人晚宴。主要出席人員17:00到場；Aimi可於18:00抵達。"),
    ("Glandys", "Glandys"),
    ("WhatsApp local contact", "WhatsApp 當地聯絡人"),
    ("Sheraton Grand Taipei", "台北喜來登大飯店"),
    ("Wang Tai Foundation", "旺台基金會"),
    ("Taipei Concert Hall", "國家音樂廳"),
    ("Taipei Recital Hall", "國家演奏廳"),
    ("Weiwuying Concert Hall", "衛武營音樂廳"),
    ("Taichung Mid Theater", "臺中國家歌劇院中劇院"),
    ("Taichung Grand Theater", "臺中國家歌劇院大劇院"),
    ("Taipei Chamber Singers Rehearsal room", "台北室內合唱團大排練室"),
    ("Steinway center", "史坦威中心"),
    ("Arrival:", "抵達："),
    ("Departure:", "離境："),
    ("Terminal", "航廈"),
    ("to Taipei", "前往台北"),
    ("to Tokyo Haneda", "前往東京羽田"),
    ("to Osaka Kansai", "前往大阪關西"),
    ("to Seoul Incheon", "前往首爾仁川"),
    ("to Beijing", "前往北京"),
    ("with", "與"),
]


def translate(text):
    if not text:
        return text
    if text in EXACT:
        return EXACT[text]
    translated = text
    for old, new in REPLACEMENTS:
        translated = translated.replace(old, new)
    translated = re.sub(r"Sep (\d+), 2026", r"2026年9月\1日", translated)
    translated = translated.replace("Sep ", "9月")
    translated = translated.replace("Morning", "上午")
    translated = translated.replace("Afternoon", "下午")
    translated = translated.replace("Evening", "晚上")
    translated = translated.replace("Confirm ", "確認")
    return translated


def replace_paragraph(paragraph, translated):
    if paragraph.text == translated:
        return
    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        return
    text_runs[0].text = translated
    for run in text_runs[1:]:
        run.text = ""


document = Document(PATH)

for paragraph in document.paragraphs:
    replace_paragraph(paragraph, translate(paragraph.text))

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_paragraph(paragraph, translate(paragraph.text))

for section in document.sections:
    for paragraph in section.header.paragraphs:
        replace_paragraph(paragraph, translate(paragraph.text.replace("MASTER OPERATIONS HANDBOOK", "總務運作手冊")))
    for paragraph in section.footer.paragraphs:
        replace_paragraph(paragraph, translate(paragraph.text.replace("Internal working document | Confirm live transport details before each departure", "內部工作文件｜每次出發前請再次確認即時交通安排")))

document.save(PATH)
print("Translated master operations handbook into Chinese.")
