from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics')
OUTPUT = Path('/private/tmp/logistics_locations_update')

TARGETS = {
    'Aimi_Kobayashi.docx',
    'Adrien_La_Marca.docx',
    'Boris_Borgolotto.docx',
    'Brannon_Cho.docx',
    'Da_Zong_Guan_Master_Schedule.docx',
    'Edgar_Moreau.docx',
    'Eugene_Lin.docx',
    'Jinjoo_Cho.docx',
    'Kyu_Yeon_Kim.docx',
    'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx',
    'Sheraton_Front_Desk.docx',
    'Sirena_Huang.docx',
    'Steven_Lin.docx',
}

SPONSOR_DINNER_FILES = {
    'Aimi_Kobayashi.docx',
    'Da_Zong_Guan_Master_Schedule.docx',
    'Eugene_Lin.docx',
    'Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx',
    'Sheraton_Front_Desk.docx',
    'Sirena_Huang.docx',
    'Steven_Lin.docx',
}

STAFF_ENTRANCE = 'https://maps.app.goo.gl/kxPsa1KHaKpFppgi9'


def full_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return '\n'.join(parts)


def add_line(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9)


def add_reference(doc, filename):
    text = full_text(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run('Key locations & access')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(13)

    if 'CHR1' in text or 'CHR2' in text:
        add_line(doc, 'CHR rehearsal rooms (CHR1 / CHR2) — National Concert Hall Backstage, Basement Floor. Use the staff entrance: ' + STAFF_ENTRANCE + '. Please bring your ID card or passport and show it to the guard before entering.')
    if 'National Concert Hall' in text or 'National Recital Hall' in text or 'Taipei Concert Hall' in text:
        add_line(doc, 'National Concert Hall / National Recital Hall — No. 21-1, Zhongshan S. Rd., Zhongzheng Dist., Taipei City 100012, Taiwan.')
    if 'Taichung' in text:
        add_line(doc, 'National Taichung Theater — 101, Huilai Rd., Sec. 2, Xitun District, Taichung City 407025, Taiwan.')
    if 'Weiwuying' in text:
        add_line(doc, 'National Kaohsiung Center for the Arts (Weiwuying) — No. 1, Sanduo 1st Rd., Fengshan Dist., Kaohsiung City 830043, Taiwan.')
    if filename in SPONSOR_DINNER_FILES:
        add_line(doc, "Sponsor's Dinner — 1F., No. 70-1, Chengde Rd., Sec. 1, Datong Dist., Taipei City, Taiwan.")


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for name in sorted(TARGETS):
        doc = Document(SOURCE / name)
        add_reference(doc, name)
        doc.save(OUTPUT / name)


if __name__ == '__main__':
    main()
