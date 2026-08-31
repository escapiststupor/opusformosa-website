from pathlib import Path

from docx import Document


SOURCE = Path('/Users/pyen/OpusFormosa/festival_planning/logistics/Da_Zong_Guan_Master_Schedule.docx')
OUTPUT = Path('/private/tmp/da_zong_guan_availability_update/Da_Zong_Guan_Master_Schedule.docx')


def main():
    doc = Document(SOURCE)
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith('You are UNAVAILABLE:'):
            paragraph.text = 'Availability: You are available throughout every festival day and should remain on call for all commitments.'
            break
    else:
        raise ValueError('Availability note not found')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'unavailable noon–afternoon' in cell.text.lower():
                    cell.text = 'Confirm car for all five → Taipei Recital Hall (Dress 09:30). Remain available throughout the day.'
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == '__main__':
    main()
