from copy import deepcopy
from docx import Document

PATH = "/Users/pyen/OpusFormosa/festival_planning/logistics/Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"

def activity_row(table, activity):
    for row in table.rows[1:]:
        if row.cells[2].text == activity:
            return row
    raise ValueError(activity)

def set_row(row, values):
    for cell, value in zip(row.cells, values):
        cell.text = value

def add_after(table, row, values):
    position = list(table._tbl).index(row._tr)
    xml_row = deepcopy(row._tr)
    table._tbl.insert(position + 1, xml_row)
    new_row = next(item for item in table.rows if item._tr is xml_row)
    set_row(new_row, values)
    return new_row

document = Document(PATH)

# Sep 3
table = document.tables[5]
schumann = activity_row(table, "Schumann: Piano Quartet Op.47")
add_after(table, schumann, ["13:00–15:00", "Adrien La Marca; Eugene Lin; Sirena Huang; Steven Lin", "Schumann: Piano Quartet Op.47", "CHR3", ""])

# Sep 7
table = document.tables[9]
set_row(activity_row(table, "Mahler: Piano Quartet in A minor"), ["13:00–15:00", "Adrien La Marca; Eugene Lin; Kyu Yeon Kim; Sirena Huang", "Mahler: Piano Quartet in A minor", "CHR4", ""])
activity_row(table, "琴房：平台鋼琴").cells[0].text = "16:00–19:00"

# Sep 8
table = document.tables[10]
set_row(activity_row(table, "Ravel: Piano Trio in A minor"), ["13:00–15:00", "Brannon Cho; Jinjoo Cho; Kyu Yeon Kim", "Ravel: Piano Trio in A minor", "CHR4", ""])
rachmaninoff = activity_row(table, "Rachmaninoff: Trio élégiaque No.1")
set_row(rachmaninoff, ["15:00–17:00", "Edgar Moreau; Jinjoo Cho; Steven Lin", "Rachmaninoff: Trio élégiaque No.1", "CHR4", ""])
add_after(table, rachmaninoff, ["15:00–17:00", "Adrien La Marca; Eugene Lin; Kyu Yeon Kim; Sirena Huang", "Mahler: Piano Quartet in A minor", "CHR2", ""])
set_row(activity_row(table, "Dvořák: Piano Quintet Op.81"), ["19:00–21:30", "Adrien La Marca; Belle Ting; Boris Borgolotto; Edgar Moreau; Steven Lin", "Dvořák: Piano Quintet Op.81", "CHR4", ""])

# Sep 9
table = document.tables[11]
faure = activity_row(table, "Fauré: Piano Quartet No.1 Op.15")
add_after(table, faure, ["15:00–17:00", "Adrien La Marca; Boris Borgolotto; Brannon Cho; Kyu Yeon Kim", "Fauré: Piano Quartet No.1 Op.15", "CHR4", ""])

# Sep 10
table = document.tables[12]
dress = activity_row(table, "彩排")
add_after(table, dress, ["15:00–17:00", dress.cells[1].text, "彩排", "國家演奏廳", ""])

# Sep 11
table = document.tables[13]
tchaikovsky = activity_row(table, "Tchaikovsky: String Sextet")
players = "Adrien La Marca; Brannon Cho; Chih-Ta Chen; Edgar Moreau; Jinjoo Cho; Sirena Huang"
set_row(tchaikovsky, ["13:00–15:00", players, "Tchaikovsky: String Sextet", "CHR4", ""])
add_after(table, tchaikovsky, ["15:00–17:00", players, "Tchaikovsky: String Sextet", "CHR4", ""])
set_row(activity_row(table, "Granados: Piano Quintet in G minor"), ["19:30–21:30", "Boris Borgolotto; Chih-Ta Chen; Edgar Moreau; Sophie Wang; Steven Lin", "Granados: Piano Quintet in G minor", "CHR3", ""])

# Sep 13
table = document.tables[15]
set_row(activity_row(table, "Brahms: String Sextet No.1 Op.18"), ["19:30–21:30", "Boris Borgolotto; Canglah Micyang; Chih-Ta Chen; Edgar Moreau; Eugene Lin; Sirena Huang", "Brahms: String Sextet No.1 Op.18", "CHR4", ""])

# Sep 14
table = document.tables[16]
dress = activity_row(table, "彩排")
performers = "Belle Ting; Boris Borgolotto; Canglah Micyang; Chih-Ta Chen; Edgar Moreau; Eugene Lin; Sirena Huang; Sophie Wang; Steven Lin"
set_row(dress, ["13:00–15:00", performers, "彩排", "臺中國家歌劇院中劇院", ""])
add_after(table, dress, ["15:00–17:00", performers, "彩排", "臺中國家歌劇院中劇院", ""])

# Sep 15
table = document.tables[17]
rachmaninoff = activity_row(table, "Rachmaninoff: Trio élégiaque No.1 (9月16 cast)")
add_after(table, rachmaninoff, ["15:00–17:00", "Belle Ting; Steven Lin", "Korngold: Much Ado About Nothing (9月16 cast)", "CHR3", ""])

document.save(PATH)
print("Reconciled daily run sheet with confirmed rehearsal entries.")
