from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path("/Users/pyen/OpusFormosa/festival_planning/logistics")


def append_note(cell, note):
    if note not in cell.text:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(note)
        run.font.name = "PingFang TC"
        run.font.size = Pt(8)


# Main handbook: retain the flight, remove the operational airport-drop-off task.
path = ROOT / "Opus_Formosa_2026_Da_Zong_Guan_Handbook.docx"
document = Document(path)
for table_index in (2, 8):
    for row in document.tables[table_index].rows[1:]:
        if "Aimi Kobayashi" in row.cells[1].text and "NH852" in " | ".join(cell.text for cell in row.cells):
            row.cells[-1].text = "不安排送機（自行前往機場）"
document.save(path)

# Aimi's own itinerary: make the self-arranged airport transfer explicit.
path = ROOT / "Aimi_Kobayashi.docx"
document = Document(path)
schedule = document.tables[1]
row = schedule.add_row()
for cell, value in zip(row.cells, ["Sep 6", "13:30", "Departure — NH852; airport transfer self-arranged", "Taipei Songshan Airport, Terminal 1"]):
    cell.text = value
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "PingFang TC"
            run.font.size = Pt(8)
document.save(path)

# Hotel and legacy master reference: explicitly state that no transfer is to be arranged.
for filename in ("Sheraton_Front_Desk.docx", "Da_Zong_Guan_Master_Schedule.docx"):
    path = ROOT / filename
    document = Document(path)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "NH852" in cell.text:
                    append_note(cell, "No airport transfer arranged / 不安排送機")
    document.save(path)

print("Removed Aimi airport-drop-off instructions and added self-arranged departure notes.")
